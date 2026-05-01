from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "PROJECT_X_Academic_Report.docx"

TITLE = "AI-Based Career Counselling, Skill Gap Analyzer and Facial Recognition-Based Interview System using Machine Learning and NLP"
SHORT_TITLE = "PROJECT-X"

TEAM = [
    {"name": "Anant Kaushik", "roll": "2218345"},
    {"name": "Sandeep Rudola", "roll": "[Fill Roll No.]"},
    {"name": "Vipul Chauhan", "roll": "[Fill Roll No.]"},
]

GUIDE_NAME = "Mr. Amit Gupta"
GUIDE_DESIGNATION = "Professor"
GROUP_NO = "486"
UNIVERSITY = "Graphic Era Hill University, Dehradun"
DEPARTMENT = "Department of Computer Science and Engineering"


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1)
    set_page_number(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Subtitle"].font.size = Pt(12)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].font.bold = False


def add_paragraph(doc: Document, text: str, *, bold: bool = False, center: bool = False, first_line: bool = True, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_line and not center:
        p.paragraph_format.first_line_indent = Inches(0.2)
    run = p.add_run(text.strip())
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        rendered = text.upper()
    elif level == 2:
        rendered = text.upper()
    else:
        rendered = text
    run = p.add_run(rendered)
    run.bold = True if level in (1, 2) else False
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(f"• {item}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_code_block(doc: Document, title: str, code: str, explanation: str) -> None:
    add_heading(doc, title, level=3)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    for line in code.rstrip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        r.font.size = Pt(9.5)
    add_paragraph(doc, explanation)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10.5)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def get_lines(path: Path, start: int, end: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    selected = lines[start - 1 : end]
    return "\n".join(selected)


def add_title_page(doc: Document) -> None:
    add_paragraph(doc, "PROJECT REPORT", bold=True, center=True, first_line=False)
    add_paragraph(doc, TITLE, bold=True, center=True, first_line=False)
    add_paragraph(doc, "Submitted in partial fulfilment of the requirement for the award of the degree of", center=True, first_line=False)
    add_paragraph(doc, "BACHELOR OF TECHNOLOGY", bold=True, center=True, first_line=False)
    add_paragraph(doc, "IN", center=True, first_line=False)
    add_paragraph(doc, "COMPUTER SCIENCE & ENGINEERING - Specialization", bold=True, center=True, first_line=False)
    add_paragraph(doc, "Submitted by", center=True, first_line=False)
    for member in TEAM:
        add_paragraph(doc, f"{member['name']}    {member['roll']}", center=True, first_line=False)
    add_paragraph(doc, "Under the guidance of", center=True, first_line=False)
    add_paragraph(doc, f"{GUIDE_NAME}", bold=True, center=True, first_line=False)
    add_paragraph(doc, GUIDE_DESIGNATION, center=True, first_line=False)
    add_paragraph(doc, f"Project Group No.: {GROUP_NO}", center=True, first_line=False)
    add_paragraph(doc, DEPARTMENT, center=True, first_line=False)
    add_paragraph(doc, UNIVERSITY, center=True, first_line=False)
    add_paragraph(doc, "May, 2026", center=True, first_line=False)
    page_break(doc)


def add_certificate(doc: Document) -> None:
    add_heading(doc, "Certificate", level=1)
    add_paragraph(
        doc,
        f"This is to certify that the project entitled \"{TITLE}\" submitted by "
        f"{', '.join(member['name'] for member in TEAM)} of Bachelor of Technology in Computer Science and Engineering - Specialization "
        f"is a bonafide record of project work carried out under my supervision in the {DEPARTMENT}, {UNIVERSITY}, during the academic session 2025-2026. "
        "The work presented in this report is original and, to the best of my knowledge, fulfils the academic requirements prescribed by the university for the final year major project."
    )
    add_paragraph(
        doc,
        "The project demonstrates a practical integration of Natural Language Processing, machine learning-based classification, interview analytics, and a web-based software architecture to address a real student employability problem. "
        "The candidates have shown satisfactory progress in requirement analysis, implementation, testing, and documentation."
    )
    add_paragraph(doc, "\n\nSignature of Guide", center=False, first_line=False)
    add_paragraph(doc, f"{GUIDE_NAME}", first_line=False)
    add_paragraph(doc, f"{GUIDE_DESIGNATION}, {DEPARTMENT}", first_line=False)
    add_paragraph(doc, "\n\nSignature of Head of Department", first_line=False)
    add_paragraph(doc, "[Name of HOD]", first_line=False)
    add_paragraph(doc, f"Head, {DEPARTMENT}", first_line=False)
    page_break(doc)


def add_declaration(doc: Document) -> None:
    add_heading(doc, "Candidate's Declaration", level=1)
    add_paragraph(
        doc,
        f"We hereby declare that the project report entitled \"{TITLE}\" submitted in partial fulfilment of the requirements for the award of the Degree of Bachelor of Technology in Computer Science and Engineering - Specialization "
        f"at {UNIVERSITY} is our original work. The report has been prepared under the supervision of {GUIDE_NAME}, {GUIDE_DESIGNATION}, and has not been submitted in part or full to any other university or institution for the award of any other degree, diploma, or certificate."
    )
    add_paragraph(
        doc,
        "Wherever the ideas, observations, design practices, and technical formulations of other authors or published sources have been used, they have been acknowledged appropriately in the references section. "
        "We also declare that the software artifacts, implementation decisions, and technical analysis described in this report genuinely correspond to the project files examined in the final codebase."
    )
    for member in TEAM:
        add_paragraph(doc, f"{member['name']}    {member['roll']}    Signature: ____________________", first_line=False)
    page_break(doc)


def add_ack(doc: Document) -> None:
    add_heading(doc, "Acknowledgement", level=1)
    paragraphs = [
        f"The completion of {SHORT_TITLE} has been possible because of the support, guidance, and encouragement we received throughout the project lifecycle. We would like to express our sincere gratitude to {GUIDE_NAME}, {GUIDE_DESIGNATION}, {DEPARTMENT}, {UNIVERSITY}, for his patient mentoring, constructive criticism, and steady support during the design, implementation, and reporting stages of this work.",
        "We are equally thankful to the faculty members of the Department of Computer Science and Engineering for building the academic foundation on which this project stands. Their lectures, lab sessions, and project reviews helped us convert a broad idea into a working software system that combines machine learning, natural language processing, and web development in a meaningful way.",
        "We also acknowledge the institution for providing the academic environment, laboratory access, and learning resources required for carrying out this project. The availability of computing facilities, internet access, and technical support significantly eased the experimentation and implementation work.",
        "A special note of thanks is due to our classmates and well-wishers who offered informal feedback while we were iterating on the user interface, refining use cases, and validating practical scenarios from the student perspective. Their comments helped us keep the system grounded in real user expectations rather than treating it as a purely theoretical exercise.",
        "Finally, we are deeply grateful to our parents and family members for their constant encouragement, patience, and faith in our efforts. Their support gave us the confidence to remain consistent through the development process and to complete this report with care and sincerity.",
    ]
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)
    page_break(doc)


def add_abstract(doc: Document) -> None:
    add_heading(doc, "Abstract", level=1)
    text = [
        f"{SHORT_TITLE} is a full-stack academic project developed to support data-driven career guidance for university students by combining resume understanding, skill-gap analysis, interview evaluation, and mentor-style feedback within a single web-based system. The project addresses a practical problem often observed in higher education: students possess fragmented knowledge about their strengths, but they do not always understand how their current profile aligns with industry roles, what skills are missing, or how their communication style affects employability.",
        "The backend of the system is implemented using Flask and exposes REST APIs for profile analysis, mock interview processing, and mentor interaction. The frontend is developed in React and provides four user-facing views: profile analysis, mock interview, mentor assistance, and a university-facing dashboard. The project uses Natural Language Processing techniques to extract email, summary cues, and skill keywords from user-provided resume text. For career recommendation, it applies TF-IDF feature extraction and a One-vs-Rest Logistic Regression classifier trained on domain-labelled sample profiles. A rule-based shortcut layer is also introduced for identifiable profiles such as digital marketing, frontend development, and AI/ML, improving interpretability for smaller datasets.",
        "One of the distinctive parts of the project is the interview analytics module. A recorded video interview is processed through Whisper for speech transcription and DeepFace for frame-level emotion analysis. The transcript is further examined to identify filler words and repetition patterns, while sentiment and communication clarity are estimated using TextBlob-driven sentiment signals and a lightweight scoring rule. These outputs are merged with the profile match score to derive an employability score that communicates readiness in a simple form that students can understand.",
        "Beyond prediction, the project attempts to be advisory in nature. A virtual mentor response module uses a generative AI service to answer student questions in the context of predicted domain, employability score, and identified missing skills. The system therefore moves beyond static classification and supports an action-oriented learning loop: understand profile, identify gaps, practice interview behaviour, and receive guidance.",
        "This report documents the problem background, literature context, requirements, architecture, design decisions, implementation flow, testing strategy, observations, limitations, and future scope of the project. The document also includes representative code snippets and output descriptions so that the technical and academic understanding of the project remains balanced."
    ]
    for paragraph in text:
        add_paragraph(doc, paragraph)
    page_break(doc)


def add_toc_and_lists(doc: Document) -> None:
    add_heading(doc, "Table of Contents", level=1)
    add_paragraph(doc, "The final page numbers for the table of contents, list of tables, and list of figures can be updated automatically in Microsoft Word after opening this document and refreshing all fields. The chapter and section structure has been prepared in a format suitable for that purpose.", first_line=False, italic=True)
    toc_items = [
        "Certificate",
        "Candidate's Declaration",
        "Acknowledgement",
        "Abstract",
        "List of Tables",
        "List of Figures",
        "Abbreviations",
        "Notations",
        "Chapter 1 Introduction and Motivation",
        "Chapter 2 Literature Review and Objectives",
        "Chapter 3 Requirement Analysis and System Analysis",
        "Chapter 4 System Design and Methodology",
        "Chapter 5 Technologies Used, Implementation, and Workflow",
        "Chapter 6 Testing, Results, Discussion, Challenges, Future Scope, and Conclusion",
        "References",
        "Appendix A Selected Code Snippets",
        "Appendix B Output Screens and Snapshot Descriptions",
    ]
    add_bullets(doc, toc_items)
    page_break(doc)

    add_heading(doc, "List of Tables", level=1)
    table_entries = [
        "Table 3.1 Functional requirements of PROJECT-X",
        "Table 3.2 Non-functional requirements of PROJECT-X",
        "Table 3.3 Hardware and software requirements",
        "Table 4.1 Core modules and responsibilities",
        "Table 5.1 Technology stack used in implementation",
        "Table 6.1 Test cases and observed results",
    ]
    add_bullets(doc, table_entries)
    page_break(doc)

    add_heading(doc, "List of Figures", level=1)
    figure_entries = [
        "Figure 3.1 High-level problem context diagram",
        "Figure 3.2 Existing versus proposed system comparison",
        "Figure 4.1 Overall system architecture",
        "Figure 4.2 Data flow for profile analysis",
        "Figure 4.3 Interview analytics processing flow",
        "Figure 5.1 User interface navigation map",
        "Figure 5.2 Mock interview recording flow",
        "Figure 6.1 Consolidated output interpretation model",
        "Figure B.1 Profile analysis screen placeholder",
        "Figure B.2 Skill gap analysis result placeholder",
        "Figure B.3 Live mock interview screen placeholder",
        "Figure B.4 Interview feedback screen placeholder",
        "Figure B.5 Mentor chat screen placeholder",
        "Figure B.6 University dashboard screen placeholder",
    ]
    add_bullets(doc, figure_entries)
    page_break(doc)

    add_heading(doc, "Abbreviations", level=1)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["AI", "Artificial Intelligence"],
            ["ML", "Machine Learning"],
            ["NLP", "Natural Language Processing"],
            ["API", "Application Programming Interface"],
            ["UI", "User Interface"],
            ["UX", "User Experience"],
            ["TF-IDF", "Term Frequency - Inverse Document Frequency"],
            ["OvR", "One-vs-Rest"],
            ["PDF", "Portable Document Format"],
            ["SRS", "Software Requirement Specification"],
            ["ASR", "Automatic Speech Recognition"],
            ["CRA", "Create React App"],
        ],
    )
    page_break(doc)

    add_heading(doc, "Notations", level=1)
    add_table(
        doc,
        ["Symbol / Term", "Meaning in this project"],
        [
            ["score", "Numeric assessment value calculated by a module"],
            ["P(domain|text)", "Classifier confidence for a domain label"],
            ["required_skills", "Reference skill list for a target role"],
            ["matched_skills", "Intersection between student and target role skills"],
            ["missing_skills", "Skills expected for the role but absent in student profile"],
            ["facial_score", "Emotion-based behavioural score from interview video"],
            ["profile_match_percentage", "Percentage overlap between user skills and role requirements"],
        ],
    )
    page_break(doc)


def add_caption_bank(doc: Document) -> None:
    add_heading(doc, "Appendix C", level=1)
    add_heading(doc, "Chapter-Wise Figure Captions for Final Screenshot Insertion", level=1)
    add_paragraph(
        doc,
        "This appendix has been added for final report preparation. Each caption is written in a submission-ready tone so that screenshots and diagrams can be inserted later without rewriting figure descriptions. The numbering follows the college instruction pattern of chapter-based figure labels.",
        first_line=False,
    )

    captions = [
        ("Figure 3.1", "Problem context of PROJECT-X showing the gap between student profile awareness, industry expectations, and interview readiness."),
        ("Figure 3.2", "Comparison between the existing fragmented guidance process and the proposed integrated career-support platform."),
        ("Figure 4.1", "Overall architecture of PROJECT-X illustrating the React frontend, Flask API layer, processing modules, and supporting infrastructure."),
        ("Figure 4.2", "Data flow for profile analysis from resume text input to domain prediction and skill-gap output."),
        ("Figure 4.3", "Interview analytics flow showing recording, transcription, facial analysis, communication review, and employability scoring."),
        ("Figure 5.1", "Main navigation interface of the PROJECT-X frontend showing profile, interview, mentor, and dashboard views."),
        ("Figure 5.2", "Student profile analysis page used to submit resume text and trigger the backend analysis pipeline."),
        ("Figure 5.3", "Recommended domain and skill-gap output generated by the system after profile analysis."),
        ("Figure 5.4", "Live mock interview screen showing webcam preview, timer, and guided interview questions."),
        ("Figure 5.5", "Interview feedback screen displaying employability score, dominant emotion, and transcript issue highlighting."),
        ("Figure 5.6", "Virtual mentor chat interface showing context-aware guidance based on profile analysis."),
        ("Figure 5.7", "Prototype university dashboard presenting batch-level readiness indicators and training insights."),
        ("Figure 6.1", "Consolidated interpretation of system outputs used for student guidance and readiness discussion."),
        ("Figure B.1", "Profile analysis screen prepared for final insertion in the appendix section."),
        ("Figure B.2", "Skill-gap output screen prepared for final insertion in the appendix section."),
        ("Figure B.3", "Mock interview recording screen prepared for final insertion in the appendix section."),
        ("Figure B.4", "Interview feedback screen prepared for final insertion in the appendix section."),
        ("Figure B.5", "Mentor guidance screen prepared for final insertion in the appendix section."),
        ("Figure B.6", "University dashboard screen prepared for final insertion in the appendix section."),
    ]
    for label, text in captions:
        add_paragraph(doc, f"{label}: {text}", first_line=False)


def chapter_1(doc: Document) -> None:
    add_heading(doc, "Chapter 1", level=1)
    add_heading(doc, "Introduction and Motivation", level=1)

    add_heading(doc, "1.1 Background of the Study", level=2)
    background = [
        "Career decision-making has become more difficult in the present employment landscape because the number of technology roles has increased rapidly while the pathways leading to those roles remain unevenly understood by students. A learner may know a programming language or a tool, yet still be unsure whether that competence is sufficient for software development, data science, artificial intelligence, cybersecurity, or some adjacent discipline. The challenge is no longer lack of information alone; rather, it is the difficulty of converting scattered information into reliable personal guidance.",
        "Traditional career counselling often depends on static questionnaires, manual observation, or one-to-one sessions with mentors. Those methods can be valuable, but they are difficult to scale in institutions with large student populations. They also vary in depth and consistency depending on the experience of the advisor, the time available per student, and the quality of self-reported information. In many cases, feedback remains broad and descriptive instead of being tied to concrete evidence from a student's resume or performance.",
        "At the same time, employability has become multidimensional. Employers look beyond academic grades and expect evidence of technical competency, domain alignment, communication ability, and a capacity to learn. Students therefore need more than a statement such as 'you are suited for software development.' They need to understand why that suggestion is made, what skills support it, what skills are absent, and how interview behaviour affects the final impression.",
        "Advances in AI, machine learning, and natural language processing make it possible to automate parts of this guidance loop in a practical and explainable way. Resume text can be processed to identify skills, machine learning models can map profile content to plausible domains, and behavioural signals can be extracted from interview transcripts and facial expressions. These techniques do not replace human judgment completely, but they can provide a strong first-level analytical system that is available instantly and consistently.",
        f"{SHORT_TITLE} emerges from this context. It is not built as a generic showcase project; instead, it is designed around a clearly observed educational need: helping students translate profile evidence into direction, preparation, and confidence. The project combines a browser-based frontend, a Flask backend, lightweight NLP, a text classification pipeline, video-based interview analysis, and mentor-style response generation within a single application."
    ]
    for paragraph in background:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.2 Motivation", level=2)
    motivation = [
        "The first motivation behind the project is the mismatch between academic preparation and perceived employability. Many students finish multiple courses and mini-projects, yet they still do not know which job families align most naturally with their accumulated profile. This uncertainty leads to random upskilling, poorly targeted applications, and low interview confidence.",
        "The second motivation is the lack of immediate feedback. Students often improve only after repeated rejection or late-stage placement training. A system that can analyze profile text, identify skill gaps, and surface actionable recommendations early in the academic journey can reduce that delay significantly.",
        "A third motivation comes from the interview process itself. Students frequently prepare technical content but underestimate behavioural aspects such as filler words, repetitive phrases, facial confidence, and clarity of speech. Even a basic analytics layer that highlights these points can help them practice more deliberately.",
        "The final motivation is institutional scalability. Universities require batch-level visibility into readiness trends, dominant skill gaps, and the type of support students need. The dashboard concept included in the project reflects that wider academic use case, even though the current prototype emphasizes the student-facing flow."
    ]
    for paragraph in motivation:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.3 Problem Statement", level=2)
    problem_statement = [
        "Students and early-stage job seekers often struggle to identify a suitable career domain, understand how well their current skill set matches that domain, and improve their interview readiness using objective feedback. Existing guidance methods are usually generic, manual, or fragmented across separate tools. There is a need for a unified web-based system that can analyze resume text, predict appropriate career domains, identify missing skills, evaluate interview behaviour, and provide practical improvement guidance in a manner that is understandable and immediately useful."
    ]
    for paragraph in problem_statement:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.4 Objectives", level=2)
    add_bullets(
        doc,
        [
            "To build a web-based student career support platform using a React frontend and Flask backend.",
            "To parse resume text and extract meaningful profile information such as email, summary clues, and skills.",
            "To recommend relevant career domains using TF-IDF feature extraction and a machine learning classifier.",
            "To compare student skills against role expectations and compute a measurable profile match score.",
            "To analyze mock interview recordings through speech transcription, transcript review, and facial expression assessment.",
            "To generate an employability score by combining profile readiness and interview performance.",
            "To provide mentor-style guidance that turns analytical output into practical next steps.",
            "To document the project in an academically structured format suitable for university submission.",
        ],
    )

    add_heading(doc, "1.5 Scope of the Project", level=2)
    scope = [
        "The project scope includes resume analysis from textual input, domain recommendation, skill-gap identification, employability estimation, video-based mock interview analysis, mentor-style feedback, and static institutional insight presentation. The system is implemented as a prototype web application that demonstrates the end-to-end idea with realistic software components.",
        "The project does not attempt to provide enterprise-grade authentication, large-scale production deployment, extensive database persistence, or fully validated industry datasets. The database layer is provisioned through Docker and PostgreSQL configuration, but the current working prototype remains largely stateless for core analytical functions. This is an important academic design choice because it keeps the focus on algorithmic flow and user interaction while leaving room for future expansion."
    ]
    for paragraph in scope:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.6 Report Organization", level=2)
    organization = [
        "The remainder of this report is organized as follows. Chapter 2 presents the literature review and clarifies project objectives in relation to existing academic and technical ideas. Chapter 3 covers requirement analysis, the existing system, the proposed system, and system-level requirements. Chapter 4 discusses architecture, data flow, modules, algorithms, and design methodology. Chapter 5 explains the technologies used and the actual implementation flow based on the project codebase. Chapter 6 presents testing, observed results, discussion, challenges, future scope, and conclusion. The references section is followed by appendices containing curated code snippets and output snapshot descriptions."
    ]
    for paragraph in organization:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.7 Significance of the Project", level=2)
    significance = [
        "The significance of this project lies in the fact that it treats employability as a process rather than a single prediction outcome. A student rarely improves merely by hearing that a role is suitable or unsuitable. Improvement begins when the student sees evidence, understands the evidence, and can act on it. This report and the software together are built around that sequence.",
        "The project is also significant because it demonstrates applied interdisciplinarity. Resume analysis belongs to NLP, domain recommendation belongs to supervised machine learning, interview behaviour connects to speech and facial analytics, and the user-facing delivery belongs to full-stack software engineering. Bringing these together in a coherent undergraduate project reflects both conceptual understanding and practical implementation capability.",
        "For institutions, the significance extends beyond individual use. A future version of this system could help departments identify dominant training gaps among final-year students, prioritize workshops, and create data-backed placement preparation strategies. In that sense, the project addresses both personal guidance and institutional planning."
    ]
    for paragraph in significance:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.8 Target Users and Beneficiaries", level=2)
    beneficiaries = [
        "The primary beneficiaries of the system are undergraduate and postgraduate students who are exploring role alignment, placement preparation, or career transitions within computing and adjacent digital fields. Such users benefit directly from the profile analysis and interview evaluation cycle.",
        "A second beneficiary group includes placement coordinators, mentors, and faculty members who need a fast first-level snapshot of student readiness. Even if they continue to provide human counselling, a system like PROJECT-X can reduce routine screening effort and help them focus on higher-value mentoring conversations.",
        "A third beneficiary group consists of students who are uncertain about how to plan their upskilling efforts. Rather than consuming random online content, they can prioritize learning based on visible missing skills and domain alignment. This makes the project educationally practical, not just technically interesting."
    ]
    for paragraph in beneficiaries:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.9 Assumptions and Constraints", level=2)
    assumptions = [
        "The current project assumes that users can provide meaningful profile text in English and can record mock interview responses using a browser with media permissions enabled. It also assumes that the machine learning environment can load the required NLP and multimedia libraries locally.",
        "The system further assumes that a small curated mapping between roles and required skills is acceptable for an academic prototype. This is a reasonable assumption because the goal is to demonstrate logical and software integration rather than to claim labour-market completeness.",
        "A practical constraint is computation. Video transcription and facial analysis are heavier than simple text parsing, which means the current prototype is more suitable for individual or small-batch usage in a local development environment than for large-scale concurrent production deployment."
    ]
    for paragraph in assumptions:
        add_paragraph(doc, paragraph)

    add_heading(doc, "1.10 Academic Relevance in the Present Context", level=2)
    academic_relevance = [
        "The relevance of this project has increased in recent years because students are now exposed to a far wider range of technical roles than before, yet they often receive less personalized interpretation of where they fit. Online resources are abundant, but clarity is not. This creates a situation where students are busy learning but not always progressing strategically.",
        "Within engineering education, employability has become an explicit outcome rather than an indirect hope. Universities increasingly measure placements, internship conversion, industry engagement, and skill readiness. A project like PROJECT-X fits naturally into this academic environment because it converts employability from a vague institutional goal into a software-supported, student-visible process.",
        "The system is therefore relevant not only as a final-year project artifact, but also as a model of how academic software can address real student pain points in a way that remains technically challenging and socially useful."
    ]
    for paragraph in academic_relevance:
        add_paragraph(doc, paragraph)
    page_break(doc)


def chapter_2(doc: Document) -> None:
    add_heading(doc, "Chapter 2", level=1)
    add_heading(doc, "Literature Review and Objectives", level=1)

    add_heading(doc, "2.1 Need for Literature Context", level=2)
    add_paragraph(doc, "A literature review is important in this project because the system sits at the intersection of multiple technical areas rather than belonging to a single narrow topic. Resume parsing comes from natural language processing, career prediction relates to supervised text classification, interview analysis intersects with speech and facial analytics, and the deployment model belongs to modern full-stack application development. Reviewing these areas together helps justify the project design choices and clarifies why a hybrid approach was adopted instead of a single heavy model.")

    add_heading(doc, "2.2 Review of Career Guidance Systems", level=2)
    career_review = [
        "Digital career guidance systems generally attempt to solve one of three tasks: recommendation, matching, or advisory support. Recommendation systems try to suggest domains or jobs based on a profile. Matching systems compare profile data with predefined requirements. Advisory systems attempt to provide improvement steps, roadmaps, or mentoring cues. Most practical platforms combine these tasks, but academic prototypes often focus on only one.",
        "A recurring observation in the literature is that generic counselling systems lose value when they do not explain the recommendation. Students need more than a label; they need interpretability. This motivated the design of PROJECT-X to include explicit skill-gap output and not just domain prediction.",
        "Another trend in recent systems is the move toward lightweight, explainable text models when the dataset is small or institution-specific. While deep transformer-based systems attract attention, classical methods such as TF-IDF with Logistic Regression still remain effective for interpretable baselines, especially in academic prototypes where training data volume is limited and model transparency is important."
    ]
    for paragraph in career_review:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.3 Resume Analysis in Academic and Industry Tools", level=2)
    resume_review = [
        "Resume analysis systems typically begin with information extraction. The first stage identifies structured fields such as email, phone number, education, and experience signals. The second stage converts textual content into features for ranking, classification, or recommendation. In production environments, entity recognition, embeddings, and domain ontologies are often used. However, for student-oriented academic applications, keyword-based skill extraction remains useful because the objective is not legal-grade document parsing but guidance-oriented profiling.",
        "In the present project, this reasoning led to a hybrid parser. A rule-based regular expression identifies email addresses reliably, a controlled skill vocabulary helps with extraction of domain-relevant technologies, and spaCy is used when available to identify likely summary sentences. This keeps the parser lightweight and robust enough for demonstration without claiming unrealistic extraction accuracy.",
        "The literature also indicates that resume-derived features should be connected to outcome interpretation. A list of extracted skills has little value unless it is converted into a recommendation or a gap analysis. For that reason, the parser in PROJECT-X feeds directly into downstream modules rather than serving as a separate isolated utility."
    ]
    for paragraph in resume_review:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.4 Text Classification for Career Domain Prediction", level=2)
    classification_review = [
        "Text classification remains one of the most established applications of machine learning in document understanding. For short and medium textual profiles, TF-IDF remains a powerful representation because it captures the relative importance of terms without requiring large computational resources. Logistic Regression, especially in one-vs-rest form, is widely used in multi-label and multi-class text prediction because it handles sparse vectors efficiently and produces interpretable class scores.",
        "The codebase reflects this standard pattern. A TF-IDF vectorizer transforms the profile text into a sparse representation, the MultiLabelBinarizer prepares target labels, and OneVsRestClassifier with Logistic Regression learns the mapping from profile content to possible domains. The dataset used in the current prototype is intentionally small and curated, so the model should be understood as a proof-of-concept academic classifier rather than a production benchmark.",
        "The review also suggests that machine learning models benefit from rule-based overrides when certain domain indicators are highly distinctive. For example, terms such as React, HTML, and CSS strongly suggest frontend development, while SEO and Google Ads clearly indicate digital marketing. This is why the project includes a rule_based_domain function ahead of the classifier pipeline."
    ]
    for paragraph in classification_review:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.5 Skill Gap Analysis as Actionable Intelligence", level=2)
    skill_gap = [
        "Skill-gap analysis converts recommendation into action. If a predicted domain is shown without listing missing competencies, the student still does not know what to improve. The literature on employability systems repeatedly highlights this gap between diagnosis and intervention.",
        "PROJECT-X implements a straightforward but useful comparison model: the extracted student skills are normalized to lowercase, compared against target-role requirements, and separated into matched and missing sets. This allows the platform to compute a completeness percentage and provide a direct pathway for improvement planning. Although the mapping is based on predefined role requirements rather than live job-market mining, it remains academically valid because it demonstrates the logic of gap-driven guidance."
    ]
    for paragraph in skill_gap:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.6 Interview Analytics and Behavioural Assessment", level=2)
    interview_review = [
        "Interview readiness is not purely technical. Students often possess acceptable skill profiles but fail to communicate them with confidence. Literature on employability emphasizes communication clarity, behavioural confidence, and conversational coherence as crucial components of final hiring outcomes.",
        "The project therefore extends beyond profile analysis into mock interview analytics. Whisper is used for automatic transcription of recorded responses, enabling transcript-level inspection. The transcript is then scanned for filler words and repeated terms, producing a lightweight clarity signal. TextBlob-based sentiment interpretation adds a basic contextual indicator, while DeepFace is used on sampled video frames to estimate dominant emotional expression and a facial score. This fusion is academically meaningful because it demonstrates how multimodal input can enrich employability assessment.",
        "It is equally important to acknowledge the limitation of such behavioural interpretation. Emotion recognition is probabilistic and context-sensitive, not absolute truth. The project treats facial analysis as guidance-oriented feedback rather than a definitive measure of personality or competence."
    ]
    for paragraph in interview_review:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.7 Gap in Existing Educational Tools", level=2)
    add_paragraph(doc, "A common weakness across many student tools is fragmentation. One application helps with resume screening, another offers interview practice, and yet another provides general counselling content. Very few academic prototypes integrate all of these elements into a single, navigable interface. The present project attempts to bridge that gap by offering a unified workflow from profile analysis to mentor guidance. This integration is the central research and implementation contribution of the project at the undergraduate level.")

    add_heading(doc, "2.8 Refined Project Objectives from Literature", level=2)
    for paragraph in [
        "The literature review supports the decision to keep the current project grounded in explainable and modular methods. Instead of promising unrealistic prediction accuracy with limited data, the system focuses on traceable logic, usable feedback, and a coherent full-stack implementation. This makes the project academically defensible and practically meaningful.",
        "The final objectives refined through the literature review are therefore threefold: first, to automate first-level career guidance using resume-driven analytics; second, to provide actionable readiness feedback through skill and interview assessment; and third, to present the entire flow within a user-friendly web application that can be extended in future institutional deployments."
    ]:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.9 Comparative Discussion of Candidate Techniques", level=2)
    comparative = [
        "A comparative reading of possible techniques shows why the chosen architecture is academically appropriate. For resume understanding, a fully neural named-entity approach could extract richer semantics, but such methods require curated training data and more engineering effort. In contrast, a controlled vocabulary plus lightweight NLP is easier to justify, debug, and demonstrate in a semester-scale project.",
        "For role prediction, transformer embeddings and deep neural networks could potentially improve generalization, but they would also reduce interpretability and complicate deployment. Because the present project uses a modest role inventory and a limited training sample, TF-IDF and Logistic Regression remain more suitable. They preserve the connection between observed keywords and predicted domain labels, which matters in educational settings.",
        "For interview assessment, end-to-end behavioural modeling could include speech fluency, prosody, head pose, eye tracking, and multimodal temporal fusion. Such a system would be significantly more complex and would raise evaluation and ethics challenges. The current project instead implements a restrained analytics approach that surfaces understandable indicators without claiming psychological certainty.",
        "For user interaction, a single-page frontend with clearly separated views is preferable to a complex role-based multi-portal design because the main pedagogical aim is clarity. The literature therefore supports the use of simple, transparent, modular techniques over opaque and overengineered alternatives."
    ]
    for paragraph in comparative:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.10 Research Gap Summary", level=2)
    research_gap = [
        "The review makes a clear gap visible: many systems solve one part of the student-readiness problem, but relatively few educational prototypes combine profile analysis, skill-gap explanation, interview review, and mentor-style feedback in one coherent interface. This gap becomes especially important in university contexts where students need both assessment and direction.",
        "Another gap concerns evidence chaining. Existing tools frequently produce an output without showing how that output should influence the next student action. PROJECT-X closes that loop by feeding the prediction into gap analysis, the gap analysis into profile match, the profile match into employability score, and the employability score into mentor context.",
        "A final gap lies in pedagogical usability. Highly accurate tools are not automatically good educational tools. The report argues that projects intended for student benefit should prioritize transparency, understandable language, and actionable feedback. This gap strongly shaped the system design."
    ]
    for paragraph in research_gap:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.11 Literature Review Summary", level=2)
    summary = [
        "The literature does not point toward a single 'best' model. Instead, it points toward design trade-offs. Systems with high complexity may deliver stronger accuracy under ideal conditions, but interpretable, modular systems are more appropriate when the goal is academic demonstration and student guidance.",
        "The chosen design of PROJECT-X therefore stands on a defensible literature-informed foundation: controlled NLP for extraction, classical ML for classification, rule-based skill-gap explanation, lightweight multimodal interview feedback, and a web-based interface for accessibility."
    ]
    for paragraph in summary:
        add_paragraph(doc, paragraph)

    add_heading(doc, "2.12 Why This Review Supports the Final Design", level=2)
    why_review = [
        "The literature review does not merely provide background citations; it directly justifies the final system architecture. Because the reviewed work indicates that interpretable models remain useful for resume-like text and that educational tools benefit from visible reasoning, the selected TF-IDF, Logistic Regression, and gap-analysis approach becomes defensible.",
        "Likewise, the review of interview and behavioural analytics supports the idea that employability is partly performative and communicative. This validates the project's decision to add transcription and facial analysis even though the core problem could have been solved with profile analysis alone.",
        "In summary, the literature justifies the project's hybrid identity: part information-retrieval system, part advisory tool, part interview-practice assistant, and part educational analytics prototype."
    ]
    for paragraph in why_review:
        add_paragraph(doc, paragraph)
    page_break(doc)


def chapter_3(doc: Document) -> None:
    add_heading(doc, "Chapter 3", level=1)
    add_heading(doc, "Requirement Analysis and System Analysis", level=1)

    add_heading(doc, "3.1 Existing System", level=2)
    existing = [
        "The existing career guidance ecosystem available to students is usually distributed across separate services and manual processes. Students may rely on faculty advice, peer suggestions, placement-cell sessions, online blogs, resume scoring tools, and independent mock interview websites. Each of these can be useful in isolation, but together they create a discontinuous support experience.",
        "Manual counselling sessions are time-bound and difficult to scale. Generic online platforms often provide broad recommendations but do not analyze actual student-specific resume text. Resume tools may focus on formatting rather than domain suitability, while interview tools may not connect their observations back to a student's technical profile. As a result, the student receives pieces of advice that are not tied together logically.",
        "Another weakness of the existing system is delayed improvement. When feedback is fragmented, students struggle to prioritize what to fix first. They may learn a new framework without understanding whether it improves role match, or practice interview answers without knowing that their skill set itself is incomplete."
    ]
    for paragraph in existing:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.2 Limitations of the Existing System", level=2)
    add_bullets(
        doc,
        [
            "Lack of a unified analytical workflow from profile input to recommendation and guidance.",
            "Heavy dependence on manual counsellors and subjective interpretation.",
            "Insufficient role-specific skill-gap visibility for students.",
            "Minimal behavioural analytics during interview practice.",
            "No easy way to merge profile strength and interview performance into a single readiness measure.",
            "Poor batch-level insight for institutions that want to monitor student readiness patterns.",
        ],
    )

    add_heading(doc, "3.3 Proposed System", level=2)
    proposed = [
        f"The proposed system, {SHORT_TITLE}, is a web-based AI-assisted career support platform that combines profile parsing, career domain prediction, skill-gap analysis, interview evaluation, and mentor interaction. The frontend provides a structured interface for students, while the backend exposes analytical APIs that process resume text and interview video.",
        "The proposed workflow begins with resume or profile text input. The backend parses the content, extracts skills, predicts candidate domains, and compares the identified skills against role-specific requirements. The resulting gap analysis is shown as an interpretable profile match score. The student can then attempt a mock interview using the browser camera and microphone. The recorded response is analyzed for transcript quality and facial expression cues. Finally, the user can ask follow-up questions through a mentor chat interface that uses the analytical context already produced by the system.",
        "This proposed system improves coherence, speed, and interpretability. It shifts the student experience from disconnected advice to a guided cycle of diagnosis, reflection, and improvement."
    ]
    for paragraph in proposed:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.4 Advantages of the Proposed System", level=2)
    add_bullets(
        doc,
        [
            "Single integrated platform for multiple employability-related tasks.",
            "Quick first-level analysis without waiting for manual review.",
            "Transparent gap analysis instead of opaque recommendations.",
            "Multimodal interview assessment using transcript and facial cues.",
            "Extensible architecture with room for database, dashboards, and role expansion.",
            "Useful both for individual students and for academic departments in future versions.",
        ],
    )

    add_heading(doc, "3.5 Functional Requirements", level=2)
    add_table(
        doc,
        ["ID", "Requirement", "Description"],
        [
            ["FR1", "Profile Input", "The system shall accept resume text or profile summary from the user."],
            ["FR2", "Resume Parsing", "The system shall extract skills, email, and summary cues from input text."],
            ["FR3", "Domain Prediction", "The system shall predict one or more suitable career domains."],
            ["FR4", "Skill Gap Analysis", "The system shall compare user skills with target domain requirements."],
            ["FR5", "Mock Interview Recording", "The frontend shall capture video and audio from the browser."],
            ["FR6", "Interview Transcription", "The backend shall transcribe recorded speech for analysis."],
            ["FR7", "Facial Analysis", "The backend shall estimate dominant emotion and facial score from sampled frames."],
            ["FR8", "Communication Review", "The system shall identify filler words, repetition, and sentiment."],
            ["FR9", "Mentor Interaction", "The system shall provide context-aware mentor responses."],
            ["FR10", "Report Download", "The system shall generate a downloadable PDF interview report."],
        ],
    )

    add_heading(doc, "3.6 Non-Functional Requirements", level=2)
    add_table(
        doc,
        ["Category", "Expectation", "Project Interpretation"],
        [
            ["Performance", "Responsive operation", "Text analysis should complete within a short interactive delay for moderate inputs."],
            ["Usability", "Simple navigation", "Users should understand the four main views without training."],
            ["Maintainability", "Modular code", "Backend modules should separate parsing, ML, interview analysis, and report generation."],
            ["Scalability", "Future-ready design", "The architecture should support larger datasets and persistence later."],
            ["Portability", "Browser and desktop compatibility", "Frontend should run in common modern browsers on standard hardware."],
            ["Explainability", "Transparent outputs", "Recommendations should be supported by visible skill gaps and score components."],
            ["Reliability", "Graceful failure", "APIs should return readable error messages when required inputs are missing."],
            ["Security", "Protected secrets and safe processing", "Sensitive keys should be moved to environment variables in production."],
        ],
    )

    add_heading(doc, "3.7 Hardware and Software Requirements", level=2)
    add_table(
        doc,
        ["Type", "Requirement"],
        [
            ["Processor", "Intel i5 equivalent or higher for smooth local development"],
            ["RAM", "8 GB minimum; 16 GB preferred for video and ML dependencies"],
            ["Storage", "At least 2 GB for source code, Python packages, and runtime artifacts"],
            ["Operating System", "Windows 10/11 or equivalent development environment"],
            ["Frontend Runtime", "Node.js and npm for React application"],
            ["Backend Runtime", "Python 3.x with Flask, spaCy, scikit-learn, TextBlob, ReportLab"],
            ["Database", "PostgreSQL provisioned through Docker Compose"],
            ["Browser", "Chrome/Edge/Firefox with camera and microphone permissions"],
        ],
    )

    add_heading(doc, "3.8 Software Requirement Specification Narrative", level=2)
    srs_paragraphs = [
        "From an SRS perspective, the system behaves as an analytical web application that processes user-provided textual and multimedia input. The primary actor is the student user, while a secondary actor is the faculty or university administrator who may consult aggregate insights in an extended version of the dashboard. External system dependencies include browser media capture APIs, machine learning libraries, optional PostgreSQL storage, and a generative AI service for mentor responses.",
        "Input constraints are relatively simple. For profile analysis, valid textual content must be provided. For interview analysis, a video file must be uploaded or recorded. The backend validates the presence of required request fields before continuing. Output expectations include structured JSON containing predictions, skill gaps, scores, and feedback, as well as a downloadable PDF report for interview output.",
        "A notable design requirement in this project is interpretability. The system is not merely expected to produce a number; it must explain that number through visible intermediate outputs such as missing skills, transcript issue highlights, dominant emotion labels, and score decomposition. This requirement strongly influenced both backend API design and frontend presentation design."
    ]
    for paragraph in srs_paragraphs:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.9 Diagram Descriptions", level=2)
    diagram_descriptions = [
        "Figure 3.1 should present a problem context diagram showing the student at the center, surrounded by profile uncertainty, missing skills, interview anxiety, and lack of personalized guidance. Arrows should converge into the proposed system, which returns career suggestions, skill-gap output, interview feedback, and mentor advice.",
        "Figure 3.2 should be a comparison diagram contrasting the existing fragmented system with the proposed integrated platform. The left side can show separate boxes for manual counselling, resume tools, and interview websites, while the right side shows a single unified platform handling all major tasks in one workflow."
    ]
    for paragraph in diagram_descriptions:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.10 Feasibility Analysis", level=2)
    feasibility = [
        "Technical feasibility for the project is high because all selected frameworks and libraries are available in the Python and JavaScript ecosystems and can be executed on standard modern student hardware. Flask, React, scikit-learn, spaCy, and browser media APIs are mature technologies with ample documentation and community support.",
        "Operational feasibility is also strong because the intended users already interact with browsers, resumes, and interview practice sessions. The system does not force unfamiliar workflows; instead, it reorganizes familiar tasks into a guided digital sequence.",
        "Economic feasibility is acceptable for an academic prototype because the software stack is largely open source. The main future cost concerns would involve large-scale deployment, persistent storage, and external AI service usage for mentor responses.",
        "Schedule feasibility was achieved by prioritizing core modules first and adding higher-value extensions gradually. This reduced project risk and ensured that even an intermediate version of the system would still demonstrate meaningful academic functionality."
    ]
    for paragraph in feasibility:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.11 User Requirements Analysis", level=2)
    user_req = [
        "Student users require a system that is easy to understand, quick to use, and specific in its feedback. They do not benefit from academic jargon or unexplained scores. Therefore, the system must present outputs as recommended domain, missing skills, employability score, and practical mentor guidance.",
        "Faculty or coordinators require a system that can eventually surface aggregate patterns rather than individual raw logs. This shaped the inclusion of a dashboard concept, even though the current implementation is intentionally lightweight.",
        "From a usability standpoint, users require immediate visibility of progress. The application therefore uses loading states, result cards, progress bars, and feedback tags. These interface choices are not cosmetic alone; they support comprehension and reduce uncertainty."
    ]
    for paragraph in user_req:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.12 Use Case Narrative", level=2)
    use_case = [
        "Use Case 1: Profile guidance. The student pastes resume text, receives a recommended domain, reviews missing skills, and decides which competencies to improve before placements. This use case targets early-stage planning.",
        "Use Case 2: Interview rehearsal. The student records a mock answer, studies transcript issues and facial feedback, and repeats practice with improved awareness. This use case targets communication readiness.",
        "Use Case 3: Guidance conversation. After viewing analytical outputs, the student asks the mentor module a targeted question such as how to improve readiness for AI/ML or Software Development. This use case turns static output into a conversational support step.",
        "Use Case 4: Academic oversight. In an extended deployment, faculty review common readiness patterns in the batch and identify which workshops or interventions should be conducted. This use case gives institutional relevance to the project."
    ]
    for paragraph in use_case:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.13 Risk Considerations", level=2)
    risks = [
        "One requirement-level risk is over-reliance on limited training data. To address this, the report explicitly frames the classifier as a proof-of-concept model and emphasizes interpretability rather than benchmark competition.",
        "Another risk is behavioural misinterpretation in interview analysis. The project mitigates this by using descriptive language such as 'feedback' and 'dominant emotion' rather than making strong psychological claims.",
        "A third risk is infrastructure sensitivity: media capture, transcription, and video analysis are dependent on device capability and runtime environment. The design therefore expects graceful fallback behaviour and readable error handling.",
        "There is also a data-protection risk associated with personal resume content and recorded interviews. A full production system would require secure storage, explicit consent handling, and stricter secrets management."
    ]
    for paragraph in risks:
        add_paragraph(doc, paragraph)

    add_heading(doc, "3.14 Requirement Prioritization", level=2)
    prioritization = [
        "Not all requirements in a project have equal importance. In PROJECT-X, profile analysis, domain prediction, skill-gap output, and interview evaluation were treated as high-priority requirements because they define the identity of the system. Without these, the application would lose its central purpose.",
        "Mentor interaction and PDF report generation were medium-priority features. They significantly improve completeness and usability, but the system can still demonstrate its analytical value even without them.",
        "Dashboard-level academic analytics were treated as lower-priority but strategically important extensions. Their presence in the interface signals future institutional relevance and helps frame the project beyond individual student usage."
    ]
    for paragraph in prioritization:
        add_paragraph(doc, paragraph)
    page_break(doc)


def chapter_4(doc: Document) -> None:
    add_heading(doc, "Chapter 4", level=1)
    add_heading(doc, "System Design and Methodology", level=1)

    add_heading(doc, "4.1 Design Philosophy", level=2)
    for paragraph in [
        "The system design follows a modular full-stack approach. Instead of building one monolithic script, the project separates user interaction, API orchestration, profile parsing, machine learning, interview analytics, and PDF report generation into distinct functional units. This design improves readability, maintainability, and future extensibility.",
        "A second design principle is staged interpretation. Each module produces an intermediate output that can be understood independently. For example, resume parsing extracts skills; prediction converts text to candidate domains; gap analysis transforms the domain into missing-skill evidence; interview analysis turns raw video into facial and communication feedback. This layered structure mirrors how a human counsellor would move from observation to explanation."
    ]:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.2 Overall System Architecture", level=2)
    architecture = [
        "At the highest level, the architecture consists of a React presentation layer, a Flask application layer, a processing layer containing parsing and machine learning logic, and an auxiliary infrastructure layer containing PostgreSQL configuration and runtime dependencies. The frontend sends HTTP requests to Flask endpoints. The backend dispatches work to parser, model, and analysis modules. Responses are returned as JSON for UI rendering or as binary PDF content for download.",
        "The architecture remains service-oriented even within a small codebase. The main Flask file acts as the orchestration entry point but delegates responsibility to specialized modules such as data_processor.py, ml_models.py, interview_analyzer.py, and report_generator.py. This separation reduces coupling between concerns and makes academic explanation easier because each file corresponds to a major conceptual block."
    ]
    for paragraph in architecture:
        add_paragraph(doc, paragraph)
    add_paragraph(doc, "Figure 4.1 should show a four-layer architecture diagram: User Interface (React) -> Flask REST API Layer -> Processing Modules (Resume Parser, Classifier, Skill Gap Analyzer, Interview Analyzer, Mentor Engine, Report Generator) -> Infrastructure Layer (Python environment, optional PostgreSQL, Docker, browser media devices).", italic=True)

    add_heading(doc, "4.3 Module Description", level=2)
    add_table(
        doc,
        ["Module", "Primary Responsibility", "Main File(s)"],
        [
            ["API Orchestration", "Receives requests, validates input, coordinates modules, returns JSON", "backend/app.py"],
            ["Resume Parsing", "Extracts email, skill list, and summary clues from profile text", "backend/data_processor.py"],
            ["Career Classification", "Converts text to domains using rules and ML classification", "backend/ml_models.py"],
            ["Skill Gap Analysis", "Matches extracted skills with target role requirements", "backend/ml_models.py"],
            ["Interview Analytics", "Transcribes video, detects transcript issues, analyses facial emotion", "backend/interview_analyzer.py"],
            ["Mentor Response", "Generates mentor-style replies based on user context", "backend/ml_models.py"],
            ["Report Generator", "Creates PDF report from interview output", "backend/report_generator.py"],
            ["Frontend UI", "Collects inputs, visualizes feedback, manages view navigation", "frontend/src/App.js"],
        ],
    )

    add_heading(doc, "4.4 Data Flow Design", level=2)
    dfd = [
        "The first major data flow begins at the profile analysis screen. A user enters resume text into the interface. The frontend sends a POST request to the /api/v1/analyze_profile endpoint. The backend parses the text, extracts student skills, predicts domain labels, obtains the top target domain, compares extracted skills against a predefined requirement list, and returns a JSON response containing student data, predicted domains, skill gaps, and profile match percentage. The frontend then visualizes missing skills and the recommended domain.",
        "The second major flow begins when the user starts a mock interview. Browser media APIs capture video and audio. After recording ends, the frontend bundles the captured video into FormData and posts it to /api/v1/mock_facial_interview. The backend stores the file temporarily, transcribes speech using Whisper, performs facial analysis on sampled frames, analyzes transcript issues, and returns structured feedback including interview score, communication analysis, and facial analysis. The temporary file is deleted after processing.",
        "The third flow concerns mentor guidance. Once a profile has been analyzed, the frontend sends the user query and domain to /api/v1/mentor_chat. The backend constructs a context prompt using domain, employability score, and missing skills, then sends it to the configured generative model. The response is returned to the frontend and added to the chat history.",
        "A fourth flow supports report download. Interview result JSON is sent to the reporting endpoint, which builds a PDF using ReportLab and streams it back to the browser as an attachment."
    ]
    for paragraph in dfd:
        add_paragraph(doc, paragraph)
    add_paragraph(doc, "Figure 4.2 should show the profile analysis DFD with entities Student, Frontend, Flask API, Parser, Classifier, Skill Gap Analyzer, and Result View. Figure 4.3 should show the interview analytics DFD with entities Browser Recorder, Flask API, Whisper, DeepFace, Communication Analyzer, and Output View.", italic=True)

    add_heading(doc, "4.5 Algorithms Used", level=2)
    algorithms = [
        "Algorithm 1: Rule-based keyword domain detection. The input text is converted to lowercase and checked against curated keyword groups for digital marketing, frontend development, AI/ML, and data science. If a strong keyword group is present, the corresponding domain is returned immediately. This reduces ambiguity in obvious cases.",
        "Algorithm 2: TF-IDF based multi-label text classification. Profile text is transformed into TF-IDF vectors, training labels are binarized, and a One-vs-Rest Logistic Regression classifier is fitted. During prediction, class probabilities are sorted and thresholded. The highest probability class is always retained so that the user receives at least one interpretable recommendation.",
        "Algorithm 3: Skill-gap comparison. Required skills and extracted student skills are converted to normalized sets. Set difference yields missing skills, set intersection yields matched skills, and completeness percentage is computed as the ratio of matched to total required skills.",
        "Algorithm 4: Transcript issue highlighting. The transcript is searched for filler tokens using regular expressions and for immediate repetition patterns using word-backreference detection. Each issue is stored with type, text range, and feedback label so that the frontend can highlight problematic spans.",
        "Algorithm 5: Facial score estimation from sampled frames. The video is sampled approximately once per second. DeepFace returns emotion labels for selected frames. Positive or neutral emotions are counted, the dominant emotion is identified, and a facial score is computed as the proportion of positive/neutral frames over total analyzed frames.",
        "Algorithm 6: Composite employability score. The profile match percentage contributes 60 percent weight and interview score contributes 40 percent weight. This simple weighted model reflects the idea that strong baseline profile alignment should still matter more than a single interview attempt, while communication remains important."
    ]
    for paragraph in algorithms:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.6 Methodology Followed in Development", level=2)
    methodology = [
        "The development methodology followed an iterative build-and-review pattern rather than a rigid linear sequence. The first iteration focused on identifying the core user journey: input profile, obtain recommendation, and visualize missing skills. Once that loop was functional, the project moved to interview analytics as a second layer of value. The mentor chat and PDF reporting features were then added to improve completeness of the user experience.",
        "From a technical perspective, the methodology can be summarized in six phases: requirement understanding, module decomposition, backend prototype creation, frontend integration, feedback loop enhancement, and report/documentation preparation. This was well suited to an undergraduate project because it allowed visible progress at each stage and made debugging more manageable.",
        "Data for machine learning was handled pragmatically. A mock labelled dataset was created to demonstrate the classification flow. The team prioritized logic validation and interface integration rather than claiming industrial-scale data coverage. This is consistent with the educational objective of demonstrating how such a system can be designed and built."
    ]
    for paragraph in methodology:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.7 Design-Level Observations", level=2)
    for paragraph in [
        "An important design observation from the codebase is that the system is stronger as an integrated demonstration than as a polished production service. For instance, the database container is configured but not deeply used in the current flows, and some API routes would need restructuring before real deployment. However, these points do not weaken the academic contribution; they instead reveal the next logical engineering steps after a successful proof of concept.",
        "Another observation is that the frontend emphasizes transparency. Results are shown with cards, progress bars, highlighted transcript issues, and visible missing skill tags. This supports the project's interpretability objective and makes the outputs easier for non-expert users to understand."
    ]:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.8 Sequence-Level Design Explanation", level=2)
    sequence = [
        "If the architecture is viewed as a sequence diagram, the profile analysis operation proceeds through a clear exchange. The user sends resume text from the browser, the React client issues an HTTP request, Flask validates the body, the parser extracts structured cues, the classifier predicts domains, the skill-gap function computes alignment, and the response is returned to the frontend. Each step adds a meaningful transformation rather than duplicating work.",
        "The interview evaluation sequence is more resource-intensive. The browser records media, sends the video payload, the backend saves a temporary file, Whisper converts speech to text, DeepFace scans sampled frames, transcript review marks filler issues, score functions aggregate the output, and the frontend finally visualizes both behavioural and textual evidence. This sequence is a strong illustration of multi-stage processing in applied software design.",
        "The mentor interaction sequence is intentionally lightweight. Its value lies in context reuse. Instead of asking the student to re-explain their profile, the frontend can send the already determined domain and related context so that the mentor module can produce a more relevant answer."
    ]
    for paragraph in sequence:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.9 Data Structure and State Considerations", level=2)
    state_considerations = [
        "The backend primarily exchanges JSON objects that contain nested dictionaries and arrays. This suits the project's modular outputs: extracted skills are represented as lists, domain predictions as ordered label arrays, and transcript issues as annotated spans with type, start, and end indexes.",
        "On the frontend, state is kept simple but expressive. Separate state variables hold analysis results, interview results, chat history, and UI loading conditions. This avoids tightly coupling unrelated views and makes it possible to reset one part of the experience without losing all context.",
        "The design also illustrates a useful principle in academic software projects: choosing simple data structures often makes cross-module integration easier than premature abstraction. The current codebase uses standard dictionaries, arrays, and primitive values effectively for that reason."
    ]
    for paragraph in state_considerations:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.10 Security and Ethics by Design", level=2)
    ethics = [
        "Although security is not the main experimental contribution of the project, responsible design still matters. Personal resume content and recorded interview data are sensitive by nature. The current prototype processes them locally and does not yet implement comprehensive persistent storage, which reduces some risk, but a production deployment would still require stronger safeguards.",
        "The design should be extended with environment-based secrets management, controlled file retention, authentication, consent notices before video analysis, and clear user messaging about how feedback should be interpreted. These are not optional extras; they become core design requirements the moment the system is used beyond a classroom demonstration.",
        "Ethically, the project adopts a guidance-oriented framing. Its scores and emotional indicators are presented as developmental cues rather than judgments of worth. This framing is important because employability tools must assist students without creating deterministic or discouraging labels."
    ]
    for paragraph in ethics:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.11 Deployment View", level=2)
    deployment = [
        "In deployment terms, the application can be visualized as a browser client connected to a Python backend service over HTTP on a local network or localhost. The backend depends on a Python environment with machine learning and media-processing libraries installed. Docker Compose prepares a PostgreSQL service for persistence and future reporting features.",
        "For a scaled deployment, the frontend could be hosted as a static build, the Flask backend could run behind a production WSGI server, and media-processing tasks could be offloaded to background workers. The current prototype does not require this complexity, but its modular design leaves space for such evolution."
    ]
    for paragraph in deployment:
        add_paragraph(doc, paragraph)

    add_heading(doc, "4.12 Blueprint for Chapter Diagrams", level=2)
    blueprint = [
        "For the final printed report, the architecture diagram should use clean rectangular blocks with left-to-right data movement. The frontend block should contain the four visible views, the backend block should contain route handlers, and the processing block should include parser, classifier, skill analyzer, interview analyzer, mentor engine, and PDF generator. A distinct infrastructure block should show PostgreSQL, Docker, and media devices.",
        "The data flow diagrams should avoid excessive visual clutter. Each should use the standard pattern of external entity, process, data store, and output arrow. Since the project is academic, readability is more important than artistic complexity. Labels should match the terms already used in the report so the narrative and diagrams reinforce one another.",
        "Where a screenshot is used instead of a drawn diagram, a short note can be added below the figure to explain that the figure is a running-system capture. This is especially appropriate for frontend views such as profile analysis, interview recording, and mentor chat."
    ]
    for paragraph in blueprint:
        add_paragraph(doc, paragraph)
    page_break(doc)


def chapter_5(doc: Document) -> None:
    add_heading(doc, "Chapter 5", level=1)
    add_heading(doc, "Technologies Used, Implementation, and Workflow", level=1)

    add_heading(doc, "5.1 Technology Stack Overview", level=2)
    add_table(
        doc,
        ["Technology", "Purpose in Project", "Justification"],
        [
            ["React", "Frontend user interface", "Supports component-based interactive UI with client-side state handling."],
            ["Tailwind CSS", "Utility-based styling", "Enables rapid styling of cards, layout, and responsive UI elements."],
            ["Flask", "Backend API framework", "Lightweight and well suited for modular Python APIs."],
            ["spaCy", "Optional NLP sentence analysis", "Used for simple summary sentence detection when language model is available."],
            ["PyPDF2", "PDF text support", "Included for extracting text from PDF resumes in extended parsing workflows."],
            ["scikit-learn", "TF-IDF and Logistic Regression classification", "Provides standard, interpretable ML components for text classification."],
            ["TextBlob", "Sentiment analysis", "Lightweight sentiment cues for transcript analysis."],
            ["Whisper", "Speech transcription", "Converts recorded interview speech into text for analysis."],
            ["DeepFace", "Facial emotion analysis", "Detects dominant emotional patterns across interview frames."],
            ["ReportLab", "PDF report generation", "Creates downloadable interview reports."],
            ["Docker Compose", "PostgreSQL environment provisioning", "Prepares database layer for persistence and future extension."],
            ["PostgreSQL", "Structured data storage layer", "Configured for future scalable persistence requirements."],
        ],
    )

    add_heading(doc, "5.2 Frontend Implementation", level=2)
    frontend_paragraphs = [
        "The frontend is organized around a single main React component with multiple view-specific render blocks. Navigation is handled through a left sidebar with buttons for profile analysis, interview, mentor, and dashboard views. This design keeps the student journey explicit: first understand the profile, then practice interview, then ask mentoring questions.",
        "State variables are used to manage resume input, analysis result, interview result, chat history, loading status, current view, and error messages. This state-driven design ensures that the UI reacts to backend output without full page reloads. For instance, once profile analysis completes, the recommended domain and missing skill tags appear immediately on the same screen.",
        "The frontend also contains a custom transcript-highlighting component. Instead of simply printing the transcript returned by the backend, it reconstructs the text with highlighted spans for filler words and repetitions. This is a thoughtful design choice because it turns abstract feedback into a concrete reading experience for the student.",
        "For interview handling, the frontend uses browser media APIs. Video and audio are captured through getUserMedia, stored in recorded chunks through MediaRecorder, previewed locally after recording, and then submitted for backend analysis. A speaking indicator based on audio volume is shown during recording, adding usability feedback."
    ]
    for paragraph in frontend_paragraphs:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.3 Backend Implementation", level=2)
    backend_paragraphs = [
        "The backend is centered on Flask routes that expose feature-level APIs. The /api/v1/analyze_profile endpoint handles resume-driven guidance. It validates the request payload, parses text, predicts domains, performs skill-gap analysis, and returns a structured response. This route effectively orchestrates the core employability logic of the project.",
        "The /api/v1/mock_facial_interview endpoint handles uploaded video files. It temporarily stores the recording, runs Whisper transcription, performs facial analysis, calculates communication results, and returns interview-focused feedback. Temporary file deletion is handled in a finally block, which is good practice for resource cleanup.",
        "The /api/v1/mentor_chat endpoint demonstrates how analytical context can be connected to generative guidance. It packages the user query together with domain and score context, then requests a concise mentor-style response. The value of this endpoint lies less in the complexity of the prompt and more in the fact that it reuses earlier analytical results instead of generating disconnected generic advice.",
        "A report generation function using ReportLab builds a simple downloadable PDF from interview analysis data. This is a useful finishing feature because it converts on-screen output into a document artifact that a student could theoretically retain for revision."
    ]
    for paragraph in backend_paragraphs:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.4 Resume Parsing Implementation", level=2)
    for paragraph in [
        "The parser module follows a hybrid strategy. It begins with resilient regular-expression extraction for email detection, then scans the text against a curated skill vocabulary covering domains such as AI/ML, software development, data science, cybersecurity, frontend, and digital marketing. Multi-word phrases are matched through substring containment, while single words are checked using word boundaries to reduce false positives.",
        "When spaCy is available, the parser also attempts to identify a sentence that resembles a summary or objective statement. This is done through lightweight heuristic filtering rather than full semantic summarization. The result is not perfect, but it is academically appropriate because it demonstrates how classical NLP tools can be combined with controlled heuristics."
    ]:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.5 Career Prediction and Skill Gap Implementation", level=2)
    prediction_paragraphs = [
        "The classification module uses a small mock dataset to represent career-domain-labelled profiles. This dataset is converted into a pandas DataFrame and vectorized using TfidfVectorizer. Target domain lists are binarized using MultiLabelBinarizer, then fitted through OneVsRestClassifier wrapped around Logistic Regression. This is a sensible architecture for sparse textual classification with multi-label tendencies.",
        "Prediction begins with rule-based keyword identification. If no rule confidently applies, the trained classifier estimates probabilities for all classes, sorts them, and returns one or more labels above a modest threshold. The best label is always preserved. This makes the behaviour predictable even on limited data.",
        "Skill-gap analysis is then performed using role-specific requirement dictionaries. These dictionaries map domains such as AI/ML, Software Development, Data Science, Frontend Development, Cybersecurity, and Digital Marketing to the skills typically expected in those roles. The comparison output is later displayed to the student in the frontend as missing skill pills and a profile match score."
    ]
    for paragraph in prediction_paragraphs:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.6 Interview Analytics Implementation", level=2)
    interview_impl = [
        "The interview analysis module is the most technically ambitious part of the project because it handles multimedia input. Whisper is loaded once at module initialization and used to transcribe uploaded interview recordings. This transcript becomes the basis for communication review. The code then searches the transcript for filler words such as 'um', 'uh', and 'you know', as well as repeated adjacent words. These issues are returned with their textual positions so the frontend can visualize them.",
        "Facial analysis uses OpenCV to open the recorded video and sample frames roughly once per second. Each sampled frame is passed to DeepFace with emotion analysis enabled. Dominant emotional states are collected across frames, and a simple score is computed from the proportion of positive or neutral expressions. Although the scoring logic is simple, it is effective for providing a first-level confidence indicator.",
        "The final interview score merges communication and facial cues, while the employability score combines interview outcome with profile match percentage. This creates a multi-source readiness measure that is still easy to understand."
    ]
    for paragraph in interview_impl:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.7 End-to-End Workflow", level=2)
    workflow = [
        "Step 1: The user opens the profile section and pastes resume text or a skill summary.",
        "Step 2: The frontend sends the text to the analyze_profile API.",
        "Step 3: The backend extracts skills, predicts career domains, computes missing skills, and returns profile analysis data.",
        "Step 4: The frontend shows recommended domain, profile match bar, and missing skill tags.",
        "Step 5: The user starts a mock interview in the interview section and records a response using browser media tools.",
        "Step 6: The video is uploaded to the mock_facial_interview API for transcription and facial analysis.",
        "Step 7: The backend produces transcript issues, dominant emotion, interview score, and employability score.",
        "Step 8: The frontend displays feedback and lets the user download a PDF report.",
        "Step 9: The user may continue into the mentor section to ask for learning advice, roadmaps, or improvement suggestions.",
    ]
    add_bullets(doc, workflow)

    add_heading(doc, "5.8 Engineering Observations from the Codebase", level=2)
    for paragraph in [
        "A close reading of the project files shows that the system is feature-rich for a student project but still evolving. There are signs of rapid prototyping, such as hard-coded development URLs, a hard-coded API key in the model file, and one route definition placed under the main execution block in app.py. These are important engineering observations that should be acknowledged honestly in an academic report.",
        "At the same time, the code demonstrates strong conceptual integration. The main user flow works across multiple technical domains without collapsing into unrelated demos. Resume parsing feeds classifier output, classifier output feeds skill gap display, interview analysis generates further scoring, and mentor guidance completes the loop. That coherence is one of the main strengths of the implementation."
    ]:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.9 Detailed Technology Discussion", level=2)
    tech_discussion = [
        "React was selected because the application relies on live state transitions between analysis stages. A static multipage approach would have made it harder to preserve result context while switching from profile analysis to mentor guidance. React therefore supports both responsiveness and modularity in the UI layer.",
        "Flask was chosen because the backend required a lightweight but flexible API framework. It allows direct route definitions, easy JSON responses, and straightforward modular imports. For an academic prototype where the logic is more important than enterprise ceremony, Flask is a very practical fit.",
        "scikit-learn was used because the machine learning task is sparse text classification, a problem domain where the library is especially strong. It provides efficient vectorization and interpretable linear models without demanding large compute resources.",
        "spaCy was included for practical NLP support. Even though the parser does not depend on advanced pipelines heavily, the use of spaCy reflects a realistic approach to language processing in Python applications.",
        "Whisper and DeepFace make the project stand out from many basic student guidance tools. Their inclusion signals an attempt to handle real multimedia input rather than restricting the problem to resume text alone.",
        "ReportLab was selected because the system required a simple, code-driven PDF generation layer. It is suitable for structured output and avoids dependence on external template engines for this prototype."
    ]
    for paragraph in tech_discussion:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.10 Step-by-Step Build Narrative", level=2)
    build_narrative = [
        "The implementation began with backend proof-of-concept development. Resume parsing and domain prediction were developed first because they define the central academic problem. Once these produced stable JSON outputs, the frontend profile analysis view was added so that users could interact with the system visually.",
        "The next stage focused on interview analytics. Browser recording was implemented in the React layer, followed by backend handling of uploaded files. After transcription and facial analysis were functioning, the transcript highlighter was introduced to improve explanatory power.",
        "The mentor module was developed after the primary analytical pipeline was available. This sequence was intentional: mentor responses are only useful when grounded in real profile context. Finally, PDF report generation was added as a documentation-oriented utility that makes the software outputs feel complete.",
        "This staged build sequence reflects a disciplined undergraduate development process. Each phase added visible value while reusing the work of previous phases."
    ]
    for paragraph in build_narrative:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.11 Frontend User Experience Notes", level=2)
    ux_notes = [
        "The user interface prioritizes clarity over decorative complexity. Cards separate cognitive tasks, the sidebar limits navigation confusion, and the use of badges, progress bars, and result sections turns raw analytical values into understandable visual elements.",
        "The interview screen is especially important from a user-experience perspective because it must reassure the user during recording. The preview panel, timer, and speaking indicator all contribute to a sense of control. Without these signals, the user might be unsure whether the recording is functioning correctly.",
        "Similarly, the mentor chat interface uses message bubbles and a visible typing state to create a conversational rhythm. This helps the guidance feel more approachable than static recommendation text."
    ]
    for paragraph in ux_notes:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.12 Backend Robustness and Improvement Areas", level=2)
    robustness = [
        "The backend demonstrates a reasonable level of modularity, but there are also identifiable improvement areas. The hard-coded key in the ML module should be externalized immediately in any serious deployment. Likewise, route registration should remain outside the main execution block so that the application object behaves consistently when imported by servers or tests.",
        "The current implementation also relies on mock data and predefined role mappings. This is acceptable academically, but it would need to evolve toward configurable datasets or persistent storage in a production scenario.",
        "Despite these limitations, the backend succeeds in its main task: it converts conceptually diverse analytical operations into a clear API surface that the frontend can use effectively."
    ]
    for paragraph in robustness:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.13 Integration Flow Between Frontend and Backend", level=2)
    integration = [
        "A technically important part of the implementation is the way the frontend and backend exchange structured information. The frontend never needs to understand the internals of TF-IDF, DeepFace, or Whisper. Instead, it consumes clearly designed API responses and visualizes them in a human-friendly format. This separation keeps the user interface clean and the analytical logic reusable.",
        "For example, the transcript highlighter does not perform any NLP computation in the browser. It receives annotated span data from the backend and reconstructs the interface from those coordinates. This demonstrates a good separation of concerns: analysis happens on the server, interpretation happens in the UI.",
        "Similarly, the profile analysis view does not compute missing skills itself. It simply renders what the backend returns. This makes the architecture easier to maintain and also more academically explainable."
    ]
    for paragraph in integration:
        add_paragraph(doc, paragraph)

    add_heading(doc, "5.14 Practical Value of the Implemented Features", level=2)
    practical_value = [
        "Each feature in the project addresses a specific student pain point. The resume analysis feature reduces uncertainty about profile direction. The skill-gap feature makes learning priorities visible. The interview evaluator helps students notice behavioural habits they may ignore during self-practice. The mentor module encourages reflective improvement instead of passive score consumption.",
        "This practical alignment between features and user problems is one reason the project feels coherent. The implementation is not a random collection of AI techniques; it is a guided response to the actual workflow of placement preparation."
    ]
    for paragraph in practical_value:
        add_paragraph(doc, paragraph)
    page_break(doc)


def chapter_6(doc: Document) -> None:
    add_heading(doc, "Chapter 6", level=1)
    add_heading(doc, "Testing, Results, Discussion, Challenges, Future Scope, and Conclusion", level=1)

    add_heading(doc, "6.1 Testing Strategy", level=2)
    add_paragraph(doc, "Testing in this project was primarily functional and scenario-driven. Because the application integrates frontend behaviour, backend APIs, NLP logic, and multimedia processing, the most meaningful test method was to validate complete user flows rather than isolated mathematical outputs alone. The team therefore examined whether each major feature accepted valid input, produced expected structured output, handled obvious invalid input, and displayed results correctly in the interface.")

    add_heading(doc, "6.2 Test Cases", level=2)
    add_table(
        doc,
        ["Test ID", "Scenario", "Expected Result", "Observed Result"],
        [
            ["TC1", "Submit empty profile text", "Validation error message", "Handled with error prompting user to provide resume text"],
            ["TC2", "Submit profile containing React, HTML, CSS", "Frontend-related recommendation", "Frontend Development domain triggered through rule-based logic"],
            ["TC3", "Submit Python, TensorFlow, deep learning profile", "AI/ML recommendation", "AI/ML predicted successfully"],
            ["TC4", "Profile analysis result rendering", "Recommended domain and missing skills shown", "Rendered in frontend with progress bar and tags"],
            ["TC5", "Record and upload interview video", "Backend accepts video and returns analysis", "Video processed and structured response returned"],
            ["TC6", "Transcript with filler words", "Highlighted filler terms in output", "Frontend rendered marked transcript spans correctly"],
            ["TC7", "Mentor query after profile analysis", "Mentor responds with domain-aware advice", "Mentor chat populated with context-based response"],
            ["TC8", "Download interview report", "PDF file generated and downloaded", "Download flow implemented via backend PDF endpoint"],
        ],
    )

    add_heading(doc, "6.3 Results and Observations", level=2)
    results = [
        "The project successfully demonstrates that a unified employability support workflow can be built from lightweight but meaningful modules. Profile analysis returns understandable outputs instead of raw classifier predictions. Missing skill visualization is especially useful because it converts analysis into immediate action points.",
        "The interview module adds an important behavioural dimension. Even a basic indication of filler words, transcript repetition, and dominant facial expression makes the experience feel more practical than a text-only guidance system. This is one of the strongest experiential contributions of the project.",
        "The mentor feature extends the system from analysis to advisory support. By reusing domain context and student-specific gaps, it gives the platform a more human-facing dimension. In the current codebase, this feature depends on external API availability, but conceptually it completes the student guidance loop well.",
        "The results also reveal the prototype nature of the solution. Accuracy claims should be made carefully because the classifier is trained on a small mock dataset. Nevertheless, for academic demonstration and system integration purposes, the project performs convincingly."
    ]
    for paragraph in results:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.4 Discussion", level=2)
    discussion = [
        "From an academic standpoint, the project is valuable because it demonstrates how multiple AI-adjacent techniques can be orchestrated to solve a practical student problem. Instead of presenting machine learning as a black-box label generator, the system grounds its output in visible evidence such as skills, transcript issues, and behavioural observations.",
        "The project also illustrates an important design lesson: in educational tools, explainability often matters more than algorithmic sophistication. A simple classifier plus a transparent gap analysis can be more useful to students than an advanced opaque model that does not justify its suggestion.",
        "Another discussion point concerns multimodal feedback. The inclusion of transcript and facial cues shows how employability can be framed as a combination of technical readiness and communication behaviour. This is closer to real placement situations than profile-only systems."
    ]
    for paragraph in discussion:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.5 Challenges Faced", level=2)
    challenges = [
        "One major challenge was integrating multiple libraries with very different runtime characteristics. Traditional text processing libraries are relatively lightweight, while Whisper and DeepFace bring heavier multimedia dependencies and increased execution cost.",
        "Another challenge was balancing realism and feasibility. A production-grade career prediction system would require a much larger labelled dataset and cleaner persistence architecture. For an academic project timeline, the team had to choose a manageable scope without losing conceptual value.",
        "Browser-based interview recording also introduces practical concerns such as media permissions, audio quality, lighting conditions, and inconsistent device hardware. These factors influence facial and transcript results and therefore had to be acknowledged in the design.",
        "A further challenge was maintaining coherent user experience while combining profile analysis, interview analytics, and mentor interaction. The frontend needed to remain simple even though the backend logic became increasingly diverse."
    ]
    for paragraph in challenges:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.6 Limitations", level=2)
    limitations = [
        "The current classifier uses a small mock dataset and therefore should not be treated as statistically validated for real-world deployment.",
        "Role requirements are predefined rather than mined from live labour-market data.",
        "Database persistence is provisioned but not deeply integrated in the present prototype.",
        "Emotion analysis and sentiment interpretation provide indicative feedback only and should not be over-interpreted.",
        "Some backend routes and security practices require further refinement before production use, especially around API key handling and route organization.",
    ]
    add_bullets(doc, limitations)

    add_heading(doc, "6.7 Future Scope", level=2)
    future_scope = [
        "The first and most obvious extension is the use of a larger and more realistic labelled dataset for career prediction. With more training data, the system could support finer-grained roles, better confidence calibration, and improved generalization.",
        "A second major enhancement would be deep database integration. Student sessions, historical scores, attempted interviews, and personalized roadmaps could be stored in PostgreSQL, enabling progress tracking over time.",
        "A third extension would involve live job-market integration. Instead of relying only on predefined skill dictionaries, the system could fetch role requirements from job portals or institutional placement data and continuously update missing skill recommendations.",
        "The interview module could also be strengthened through question-wise evaluation, speech-rate analysis, pause detection, and richer behavioural dashboards. Similarly, the mentor system could evolve into a roadmap generator that produces week-wise learning plans.",
        "Finally, the university dashboard could be transformed into a genuine administrative analytics system with batch trends, at-risk student detection, workshop recommendations, and departmental employability reports."
    ]
    for paragraph in future_scope:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.8 Conclusion", level=2)
    conclusion = [
        f"{SHORT_TITLE} successfully demonstrates the design and implementation of an AI-assisted career support platform that is grounded in a real educational problem. The system integrates resume understanding, career recommendation, skill-gap analysis, interview review, mentor-style advice, and report generation into a single full-stack application.",
        "The project shows that even with lightweight and interpretable models, it is possible to create a system that provides practical value to students. By combining evidence from textual profile analysis and interview behaviour, the application moves beyond static recommendation and supports a more complete employability narrative.",
        "As an undergraduate academic project, the work is strongest in its integration, interpretability, and relevance. Its limitations are visible and honest, but they also reveal a strong future path for further development. Overall, the project achieves its primary aim of bridging the gap between academic profile awareness and career readiness through a technically meaningful and user-oriented solution."
    ]
    for paragraph in conclusion:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.9 Lessons Learned", level=2)
    lessons = [
        "One of the strongest lessons from the project is that successful student software projects do not require the most advanced algorithm in every module. They require a coherent problem statement, a careful choice of methods, and enough integration that the final system feels purposeful. PROJECT-X illustrates this well.",
        "A second lesson is that user-facing transparency matters. The parts of the system that feel most valuable are those that explain themselves well: missing skills, highlighted transcript issues, readable scores, and mentor advice grounded in visible context.",
        "A third lesson concerns engineering maturity. Real systems demand not only functionality but also configuration hygiene, data protection, robust deployment practices, and maintainable code structure. The project therefore served not only as an AI exercise but also as a software engineering learning experience."
    ]
    for paragraph in lessons:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.10 Ethical Reflection", level=2)
    ethical_reflection = [
        "Any system that analyzes personal profiles and interview recordings should be used responsibly. Scores can influence confidence, and behavioural feedback can be misread if presented carelessly. The present project therefore treats all outputs as advisory indicators rather than definitive judgments about a student's capability.",
        "This ethical stance is especially important in educational environments. Tools should help students grow, not reduce them to labels. A supportive interface, transparent explanation, and honest limitations are therefore just as important as the underlying algorithms.",
        "Future deployments should include explicit user consent, retention controls for uploaded media, and carefully worded institutional guidelines about how the outputs are to be used."
    ]
    for paragraph in ethical_reflection:
        add_paragraph(doc, paragraph)

    add_heading(doc, "6.11 Final Discussion on Readiness for Submission", level=2)
    readiness = [
        "From a submission perspective, the project is strong enough to support a detailed academic report because it contains real source code, a defined architecture, visible outputs, meaningful algorithms, and a defensible practical objective. It is neither a purely conceptual proposal nor a single-file demonstration. That gives the report substantial academic weight.",
        "Before final printing, the document should still be polished through three practical steps: insertion of actual screenshots, refresh of table-of-contents fields and page numbers in Microsoft Word, and replacement of remaining identity placeholders such as missing roll numbers or HOD details.",
        "Once those steps are completed, the report will read more naturally as a finished student submission rather than as a draft assembled from mixed sources."
    ]
    for paragraph in readiness:
        add_paragraph(doc, paragraph)
    page_break(doc)


def references_section(doc: Document) -> None:
    add_heading(doc, "References", level=1)
    refs = [
        "Manning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.",
        "Mitchell, T. M. (1997). Machine Learning. McGraw-Hill.",
        "Jurafsky, D., & Martin, J. H. (2025). Speech and Language Processing (3rd ed. draft). Stanford University.",
        "Geron, A. (2019). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (2nd ed.). O'Reilly Media.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Van Rossum, G., & Drake, F. L. (2009). Python 3 Reference Manual. CreateSpace.",
        "The Pallets Projects. (2026). Flask Documentation (3.1.x). Retrieved April 24, 2026, from https://flask.palletsprojects.com/",
        "React Team. (2026). React Documentation. Retrieved April 24, 2026, from https://react.dev/",
        "spaCy Authors. (2026). spaCy Usage Documentation: Models and Languages. Retrieved April 24, 2026, from https://spacy.io/usage/models",
        "Scikit-learn Developers. (2026). TfidfVectorizer Documentation. Retrieved April 24, 2026, from https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
        "Scikit-learn Developers. (2026). LogisticRegression Documentation. Retrieved April 24, 2026, from https://scikit-learn.org/1.6/modules/generated/sklearn.linear_model.LogisticRegression.html",
        "OpenAI. (2022, September 21). Introducing Whisper. Retrieved April 24, 2026, from https://openai.com/index/whisper/",
        "OpenAI. (2026). Whisper GitHub Repository. Retrieved April 24, 2026, from https://github.com/openai/whisper",
        "Loria, S. (2026). TextBlob Documentation (v0.19.0). Retrieved April 24, 2026, from https://textblob.readthedocs.io/",
        "Serengil, S. I. (2026). DeepFace GitHub Repository. Retrieved April 24, 2026, from https://github.com/serengil/deepface",
        "Serengil, S., & Ozpinar, A. (2024). A benchmark of facial recognition pipelines and co-usability performances of modules. Journal of Information Technologies, 17(2), 95-107.",
        "PostgreSQL Global Development Group. (2026). PostgreSQL Documentation. Retrieved April 24, 2026, from https://www.postgresql.org/docs/",
        "ReportLab. (2026). ReportLab Documentation. Retrieved April 24, 2026, from https://docs.reportlab.com/",
        "Create React App Contributors. (2026). Create React App Documentation. Retrieved April 24, 2026, from https://create-react-app.dev/",
        "Tailwind Labs. (2026). Tailwind CSS Documentation. Retrieved April 24, 2026, from https://tailwindcss.com/docs",
        "OpenCV Team. (2026). OpenCV Documentation. Retrieved April 24, 2026, from https://docs.opencv.org/",
        "Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 9(3), 90-95.",
        "McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference, 51-56.",
        "Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513-523.",
        "Turnitin. (2026). Turnitin and plagiarism. Retrieved April 24, 2026, from https://guides.turnitin.com/hc/en-us/articles/34400565079053-Turnitin-and-plagiarism",
        "Turnitin. (2026). Using the AI Writing Report. Retrieved April 24, 2026, from https://guides.turnitin.com/hc/en-us/articles/22774058814093-Using-the-AI-Writing-Report",
        "GPTZero. (2026). AI Content Detector. Retrieved April 24, 2026, from https://www.gptzero.app/",
        "Copyleaks. (2026). AI Detector. Retrieved April 24, 2026, from https://copyleaks.com/ai-content-detector/",
        "Copyleaks. (2026). Plagiarism Checker. Retrieved April 24, 2026, from https://copyleaks.com/plagiarism-checker",
        "QuillBot. (2026). Plagiarism Checker. Retrieved April 24, 2026, from https://quillbot.com/plagiarism-checker",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)
    page_break(doc)


def appendix_code(doc: Document) -> None:
    add_heading(doc, "Appendix A", level=1)
    add_heading(doc, "Selected Code Snippets and Explanations", level=1)
    add_paragraph(doc, "As required by the college format, code is included in the appendix rather than in the main theoretical chapters. Only representative and high-value snippets have been selected. These snippets show how the major ideas of the project are implemented without unnecessarily dumping entire files.", first_line=False)

    snippets = [
        (
            "A.1 Flask Profile Analysis Route",
            get_lines(ROOT / "backend" / "app.py", 59, 98),
            "This route is the main entry point for resume-driven guidance. It validates the request body, parses the resume text, predicts career domains, computes the skill-gap analysis for the highest-ranked domain, and returns all results in a structured JSON response. Academically, this snippet is significant because it demonstrates service orchestration: several modules are invoked but the API still returns a unified result for the frontend."
        ),
        (
            "A.2 Resume Parsing Logic",
            get_lines(ROOT / "backend" / "data_processor.py", 23, 60),
            "This snippet shows the lightweight resume parsing strategy. The implementation uses regular expressions for email detection, a curated vocabulary for skill extraction, and optional spaCy support for finding a summary-like sentence. It is representative of a practical academic approach: simple techniques are used where they are reliable, instead of adding unnecessary complexity."
        ),
        (
            "A.3 Career Classifier and Domain Prediction",
            get_lines(ROOT / "backend" / "ml_models.py", 38, 91),
            "This section is the core machine learning logic of the project. It includes the rule-based domain detector, the TF-IDF vectorizer, the One-vs-Rest Logistic Regression classifier, training over mock data, and probability-based label selection. The snippet matters because it translates theoretical classification ideas into an interpretable implementation."
        ),
        (
            "A.4 Skill Gap Analysis and Mentor Response",
            get_lines(ROOT / "backend" / "ml_models.py", 108, 154),
            "The first part of this snippet compares required skills with student skills and computes completeness percentage. The second part shows the mentor response prompt construction using domain, employability score, and missing skills. Together, these functions illustrate how the system moves from diagnosis to recommendation."
        ),
        (
            "A.5 Interview Analytics Pipeline",
            get_lines(ROOT / "backend" / "interview_analyzer.py", 10, 129),
            "This is one of the most meaningful snippets in the entire codebase because it handles multimodal input. It loads Whisper, transcribes the interview, samples video frames for emotion analysis through DeepFace, highlights transcript issues, and computes communication and employability-related scores. This snippet embodies the project's claim of combining NLP and behavioural analysis."
        ),
        (
            "A.6 Transcript Highlighter and Interview UI Flow",
            get_lines(ROOT / "frontend" / "src" / "App.js", 14, 65),
            "This frontend snippet shows how transcript issues are reconstructed and visually highlighted. Its importance is not algorithmic but experiential: it turns backend analytics into a readable interface that students can immediately act upon."
        ),
        (
            "A.7 Interview Recording and Evaluation Flow",
            get_lines(ROOT / "frontend" / "src" / "App.js", 157, 312),
            "This larger snippet demonstrates browser media handling, recording control, upload flow, and the display of evaluation results. It is representative of the bridge between user interaction and backend multimedia analytics."
        ),
        (
            "A.8 PDF Report Generation",
            get_lines(ROOT / "backend" / "report_generator.py", 1, 36),
            "This code creates a downloadable PDF summarizing the interview output. It shows how analysis results can be materialized into a document, which is useful for review, sharing, or placement preparation."
        ),
    ]
    for title, code, explanation in snippets:
        add_code_block(doc, title, code, explanation)
    page_break(doc)


def appendix_output(doc: Document) -> None:
    add_heading(doc, "Appendix B", level=1)
    add_heading(doc, "Output Screens and Snapshot Descriptions", level=1)
    add_paragraph(doc, "Where running screenshots are not embedded directly into this draft, the following descriptions can be used to capture and label final snapshots from the working application. Each description is written in a report-ready manner so that screenshots can be inserted later with matching captions.", first_line=False)

    sections = [
        ("Figure B.1 Profile Analysis Screen", "This screenshot should show the profile input interface with the textarea used for pasting resume text. The left panel contains the input box and the 'Analyze Profile & Skills' button, while the right panel displays the analysis card. The caption should mention that the student enters resume or profile text to begin the employability workflow."),
        ("Figure B.2 Recommended Domain and Skill Gap Output", "This screenshot should capture the result after successful profile analysis. It should include the recommended domain, profile match progress bar, and missing skill tags. The caption should highlight that the system does not merely predict a domain but also identifies the specific competencies still required for better readiness."),
        ("Figure B.3 Live Mock Interview Recording Screen", "This image should show the camera preview, timer, speaking indicator, and list of interview questions during an active mock interview session. The caption may explain that the user records a practice response using browser camera and microphone access."),
        ("Figure B.4 Interview Performance Feedback Screen", "This screenshot should show the employability score, dominant emotion, transcript-highlighter output, and feedback text after the interview is evaluated. It is one of the most important outputs because it combines behavioural and textual indicators."),
        ("Figure B.5 Mentor Chat Interface", "This snapshot should display the mentor chat window with student and mentor messages. Ideally, it should show a query related to the recommended domain and a mentor response suggesting improvements, learning direction, or roadmap steps."),
        ("Figure B.6 University Dashboard View", "This screenshot should capture the administrative dashboard card showing average employability score, top skill gap, most recommended domain, and students requiring mentorship. The caption should clarify that this view is a prototype extension for institutional monitoring."),
        ("Figure B.7 Downloaded Interview PDF Report", "If possible, insert an image or first page preview of the downloaded PDF report. The caption should state that the system generates a portable document summarizing interview performance and key observations."),
    ]
    for title, text in sections:
        add_heading(doc, title, level=3)
        add_paragraph(doc, text)

    add_heading(doc, "Screenshot Capture Checklist", level=3)
    checklist = [
        "Capture all screenshots at the same zoom level and browser width so the report looks consistent.",
        "Use the same dataset/profile text when possible so the story across figures remains connected.",
        "Ensure the browser tab and desktop clutter are not visible in the final screenshot unless required.",
        "For interview screenshots, test camera framing and room lighting before final capture.",
        "Insert each screenshot immediately below its matching figure caption to avoid numbering confusion.",
    ]
    add_bullets(doc, checklist)

    add_heading(doc, "Suggested Screenshot Order for Final Report", level=3)
    order_text = [
        "1. Application home/profile page before analysis.",
        "2. Profile analysis result with recommended domain and missing skills.",
        "3. Live mock interview screen while recording.",
        "4. Interview feedback screen after evaluation.",
        "5. Mentor chat with at least one meaningful reply visible.",
        "6. University dashboard summary card.",
        "7. Downloaded interview PDF first page.",
    ]
    for item in order_text:
        add_paragraph(doc, item, first_line=False)


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_certificate(doc)
    add_declaration(doc)
    add_ack(doc)
    add_abstract(doc)
    add_toc_and_lists(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    references_section(doc)
    appendix_code(doc)
    appendix_output(doc)
    page_break(doc)
    add_caption_bank(doc)
    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
