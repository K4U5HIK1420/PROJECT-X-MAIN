from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "PROJECT_X_Conference_Paper_Safer.docx"
FIG = ROOT / "report_figures" / "conference_architecture_safer.png"


def set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "720")
    if not sect_pr.xpath("./w:cols"):
        sect_pr.append(cols)


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_heading_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, bold=True, size=10)


def add_body_paragraph(doc: Document, text: str, *, first_line_cm: float = 0.32) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, size=10)


def add_reference(doc: Document, idx: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"[{idx}] {text}", size=9)


def create_architecture_figure() -> None:
    FIG.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1700, 860), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("timesbd.ttf", 38)
        head_font = ImageFont.truetype("timesbd.ttf", 26)
        body_font = ImageFont.truetype("times.ttf", 22)
        small_font = ImageFont.truetype("times.ttf", 18)
    except OSError:
        title_font = ImageFont.load_default()
        head_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    def box(x1, y1, x2, y2, fill, title, items):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, outline="black", width=3, fill=fill)
        draw.text(((x1 + x2) / 2, y1 + 25), title, anchor="ma", font=head_font, fill="black")
        y = y1 + 90
        for item in items:
            draw.text((x1 + 22, y), f"- {item}", font=body_font, fill="black")
            y += 44

    draw.text((850, 45), "PROJECT-X Prototype Architecture", anchor="ma", font=title_font, fill="black")
    box(60, 140, 390, 700, "#dbeafe", "Frontend", [
        "Resume text input",
        "Mock interview screen",
        "Mentor chat view",
        "Transcript highlighter",
        "Dashboard card view",
    ])
    box(450, 140, 800, 700, "#dcfce7", "Flask Routes", [
        "/api/v1/analyze_profile",
        "/api/v1/mock_facial_interview",
        "/api/v1/mentor_chat",
        "/api/v1/interview_report",
        "JSON and file handling",
    ])
    box(860, 140, 1220, 700, "#fef3c7", "Processing Modules", [
        "simple_resume_parser",
        "CareerClassifier",
        "analyze_skill_gap",
        "transcribe_video",
        "analyze_video_faces",
    ])
    box(1280, 140, 1640, 700, "#fee2e2", "Support", [
        "Whisper base model",
        "DeepFace / fallback mode",
        "TextBlob sentiment",
        "ReportLab PDF output",
        "Docker + PostgreSQL setup",
    ])

    for x1, x2, label in ((390, 450, "request"), (800, 860, "call"), (1220, 1280, "runtime")):
        y = 410
        draw.line((x1 + 8, y, x2 - 18, y), fill="black", width=5)
        draw.polygon([(x2 - 18, y - 11), (x2 - 18, y + 11), (x2, y)], fill="black")
        draw.text(((x1 + x2) / 2, y - 28), label, anchor="ma", font=small_font, fill="black")

    img.save(FIG)


def main() -> None:
    create_architecture_figure()

    doc = Document()
    configure_page(doc.sections[0])
    set_normal_style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    add_run(
        p,
        "PROJECT-X: A Prototype Web System for Resume-Based Career Guidance and Interview Feedback",
        bold=True,
        size=16,
    )

    authors = doc.add_table(rows=1, cols=3)
    authors.alignment = WD_TABLE_ALIGNMENT.CENTER
    blocks = [
        ("Anant Kaushik", "Dept. of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
        ("Sandeep Rudola", "Dept. of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
        ("Vipul Chauhan", "Dept. of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
    ]
    for cell, (name, aff) in zip(authors.rows[0].cells, blocks):
        set_cell_borderless(cell)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cp, name + "\n", bold=True, size=10)
        add_run(cp, aff, size=9)

    abs_p = doc.add_paragraph()
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.space_after = Pt(3)
    add_run(abs_p, "Abstract—", bold=True, italic=True, size=10)
    add_run(
        abs_p,
        " In this work we describe the prototype that we built for student career guidance and interview practice. The system accepts resume text, predicts likely job domains, compares the student profile with a small skill requirement list, and then gives interview feedback from a recorded response. The backend was implemented in Flask with a lightweight resume parser, a TF-IDF plus logistic regression classifier, a rule-based shortcut for obvious domains, and separate modules for transcription and facial analysis. The interview side uses Whisper when available and falls back safely when heavier dependencies are missing. Our goal was not to claim a final recruitment product. The goal was to build one working academic prototype that joins profile analysis, skill-gap reporting, mentor chat, and interview feedback in the same flow. The paper focuses on what was actually implemented, what worked well in local testing, and which parts still remain prototype-level.",
        italic=True,
        size=10,
    )

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.paragraph_format.space_after = Pt(6)
    add_run(kw, "Keywords—", bold=True, italic=True, size=10)
    add_run(
        kw,
        "career guidance, skill gap analysis, interview feedback, Flask, React, TF-IDF, Whisper, DeepFace",
        italic=True,
        size=10,
    )

    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_page(sec)
    set_columns(sec, 2)

    add_heading_paragraph(doc, "I. INTRODUCTION")
    add_body_paragraph(
        doc,
        "Students often know some tools and languages but still do not know which role fits them best or which skills they are missing for that role. In our project work we saw this problem clearly. Resume analyzers only look at text. Interview tools mostly look at speaking practice. General chat systems can answer questions, but they do not know the student's current profile unless that context is supplied again and again.",
    )
    add_body_paragraph(
        doc,
        "Because of that, we built PROJECT-X as one combined workflow. A student can paste resume text, get a likely domain, see matched and missing skills, record a mock interview, and then ask follow-up questions in the mentor section. The main idea is simple: the output of one part of the system should be useful to the next part, instead of making the user jump between separate tools.",
    )
    add_body_paragraph(
        doc,
        "This paper is intentionally centered on the implementation rather than broad claims. We do not present a large benchmark study. What we present is the design of the prototype, the logic inside the main modules, and the observations we made while running it locally.",
    )

    add_heading_paragraph(doc, "II. IMPLEMENTED SYSTEM")
    add_body_paragraph(
        doc,
        "The backend is split across a few small files: app.py, data_processor.py, ml_models.py, interview_analyzer.py, and report_generator.py. The main routes are /api/v1/analyze_profile, /api/v1/mock_facial_interview, /api/v1/mentor_chat, and /api/v1/interview_report. On the frontend side, React is used to switch between profile analysis, interview, mentor, and dashboard views inside a single interface.",
    )
    add_body_paragraph(
        doc,
        "The resume flow starts in the analyze_profile route. The system checks whether resume_text was sent, parses the text, predicts domains, and then compares the extracted skills with the requirement list of the top predicted domain. The response contains student data, career recommendations, a skill gap object, and a profile_match_percentage value that the UI shows directly.",
    )
    add_body_paragraph(
        doc,
        "The interview flow is more practical than theoretical. The route saves the uploaded video as temp_interview.webm, sends it to the transcription and facial analysis functions, and then returns interview_score, communication_analysis, and facial_analysis in JSON. The current route still returns a fixed employability score of 75 in one place, which is a useful reminder that the project is a working prototype, not a finished product.",
    )

    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.paragraph_format.space_after = Pt(2)
    r = fig_p.add_run()
    r.add_picture(str(FIG), width=Inches(3.1))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    add_run(cap, "Fig. 1. Main blocks used in the current PROJECT-X prototype.", size=8)

    add_heading_paragraph(doc, "III. PROFILE ANALYSIS MODULE")
    add_body_paragraph(
        doc,
        "The resume parser in data_processor.py is lightweight by design. It uses regular expressions for basic extraction such as email detection and then checks a curated skill list against the resume text. If spaCy is available, it also tries to find a summary-like sentence. We kept this module simple because it was easier to debug and easier to explain in the academic setting.",
    )
    add_body_paragraph(
        doc,
        "For domain prediction, the code uses a hybrid setup. First, a rule-based check looks for obvious cases such as React and JavaScript for Frontend Development, TensorFlow and CNN for AI/ML, or SEO and Google Ads for Digital Marketing. If no clear rule matches, the text is passed through a TF-IDF vectorizer and an OneVsRestClassifier built on LogisticRegression. The training data is small: seven mock profile examples in ml_models.py. A threshold of 0.12 is used after sorting prediction probabilities, and the best domain is always kept even if only one label is strong enough.",
    )
    add_body_paragraph(
        doc,
        "The skill gap function is straight to the point. It lowers both required skills and extracted skills, computes set intersection and set difference, and returns matched_skills, missing_skills, and completeness_percentage. That percentage is later shown in the UI as the profile match value. It is a simple method, but in practice it gives students something concrete to act on.",
    )

    add_heading_paragraph(doc, "IV. INTERVIEW ANALYSIS MODULE")
    add_body_paragraph(
        doc,
        "The interview_analyzer.py file contains four useful pieces: video transcription, facial analysis, transcript issue highlighting, and communication scoring. Whisper base is loaded when available. If it is not available, the system returns a fallback transcript instead of crashing. We considered this important because heavy dependencies may not be installed on every local machine used for demonstration.",
    )
    add_body_paragraph(
        doc,
        "Facial analysis also has a fallback path. If cv2 or DeepFace is missing, the function returns a safe neutral response with a fallback message. When the full stack is available, the code samples roughly one frame per second and runs DeepFace.analyze with emotion detection. Positive emotions in the current logic are happy, neutral, and surprise. The facial_score is computed from the proportion of those frames.",
    )
    add_body_paragraph(
        doc,
        "Transcript issue detection is rule-based. The current filler list contains um, uh, umm, like, and you know. Repetition is detected with a back-reference pattern. The communication score starts from 100 and drops by 5 for each detected issue, but it is not allowed to go below 50. This is not a research-grade scoring model. It is a simple classroom-friendly scoring rule that made the feedback readable during testing.",
    )

    add_heading_paragraph(doc, "V. FRONTEND AND USER FLOW")
    add_body_paragraph(
        doc,
        "On the frontend, the most useful part is not a model but the way results are shown. App.js includes a TranscriptHighlighter component that rebuilds the transcript with colored spans around filler words and repeated terms. This matters because students can see exactly where the issue is instead of reading a vague warning like \"improve communication.\"",
    )
    add_body_paragraph(
        doc,
        "The interface also keeps the main actions visible. The user can analyze profile text, move to the mentor view, record an interview, and download a PDF report. In the mentor section the placeholder text changes depending on whether a domain has already been predicted. That small interaction made the workflow feel more connected during local runs.",
    )

    add_heading_paragraph(doc, "VI. LOCAL OBSERVATIONS")
    add_body_paragraph(
        doc,
        "During local testing, the strongest point of the project was not model accuracy. It was continuity. The same student input could move through several useful stages without resetting context. Resume analysis produced a target domain, the gap analysis showed why that domain was suggested, the interview module added communication evidence, and the mentor route could then answer a more informed question.",
    )
    add_body_paragraph(
        doc,
        "We also noticed a few prototype limitations quickly. The mock training set is too small for strong claims. The mentor component depends on an external API key and falls back to generic advice when that path fails. The facial module can only be as good as the local environment and lighting conditions. These are not minor details; they shape how the project should be presented honestly.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Part"
    hdr[1].text = "What the code currently does"
    hdr[2].text = "Why it mattered in testing"
    rows = [
        ("Domain prediction", "Rule match first, else TF-IDF + logistic regression", "Gave a quick role suggestion from resume text"),
        ("Skill gap", "Set difference on required vs extracted skills", "Turned prediction into an actionable list"),
        ("Transcript review", "Flags filler words and repetitions", "Made speaking issues visible line by line"),
        ("Facial analysis", "Emotion sampling with fallback mode", "Allowed demo even on partial setups"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    tcap = doc.add_paragraph()
    tcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tcap.paragraph_format.space_after = Pt(4)
    add_run(tcap, "Table I. A few practical behaviors of the current implementation.", size=8)

    add_heading_paragraph(doc, "VII. LIMITATIONS AND NEXT STEPS")
    add_body_paragraph(
        doc,
        "The present version should not be treated as a final employability engine. The classifier was trained on a tiny mock dataset. The scoring rules are hand-built. The mock_facial_interview route still contains a fixed employability score in one response path. Some modules also depend on optional packages, so fallback behavior is part of the design. These choices were acceptable for a project prototype, but they are not enough for deployment claims.",
    )
    add_body_paragraph(
        doc,
        "A more mature version would need a better dataset, proper evaluation metrics, stronger authentication, safer handling of uploaded media, and cleaner persistence around reports and user history. Even so, the current system already works well enough to show the basic idea in a live academic demo.",
    )

    add_heading_paragraph(doc, "VIII. CONCLUSION")
    add_body_paragraph(
        doc,
        "PROJECT-X shows that a student employability support tool can be built with simple but connected modules. In our case, the useful result came from putting the parts together: resume parsing, domain prediction, skill-gap reporting, interview feedback, and mentor guidance. The system is still rough in places, but it is real, it runs, and it gives students something more practical than a static report or a single score.",
    )

    add_heading_paragraph(doc, "ACKNOWLEDGMENT")
    add_body_paragraph(
        doc,
        "We thank the Department of Computer Science and Engineering, Graphic Era Hill University, for academic guidance and for supporting the development of this undergraduate project.",
        first_line_cm=0.0,
    )

    add_heading_paragraph(doc, "REFERENCES")
    refs = [
        "C. D. Manning, P. Raghavan, and H. Schutze, Introduction to Information Retrieval. Cambridge, U.K.: Cambridge Univ. Press, 2008.",
        "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” J. Mach. Learn. Res., vol. 12, pp. 2825-2830, 2011.",
        "D. Jurafsky and J. H. Martin, Speech and Language Processing, 3rd ed. draft. Stanford, CA, USA: Stanford Univ., 2025.",
        "A. Radford et al., “Robust speech recognition via large-scale weak supervision,” in Proc. ICML, 2023.",
        "The Pallets Projects, “Flask documentation,” 2026. [Online]. Available: https://flask.palletsprojects.com/",
        "A. Geron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 2nd ed. Sebastopol, CA, USA: O’Reilly, 2019.",
        "S. Bird, E. Klein, and E. Loper, Natural Language Processing with Python. Sebastopol, CA, USA: O’Reilly, 2009.",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)

    doc.save(OUT)
    print(f"Saved safer conference paper to: {OUT}")


if __name__ == "__main__":
    main()
