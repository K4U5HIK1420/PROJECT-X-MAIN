from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "PROJECT_X_Conference_Paper.docx"
FIG = ROOT / "report_figures" / "conference_architecture.png"


def set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_paragraph_borderless(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right", "between"):
        el = p_bdr.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            p_bdr.append(el)
        el.set(qn("w:val"), "nil")


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "720")
    if not sect_pr.xpath("./w:cols"):
        sect_pr.append(cols)


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_heading_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, bold=True, size=10)


def add_body_paragraph(doc: Document, text: str, *, first_line_cm: float = 0.32) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, size=10)


def add_reference(doc: Document, idx: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"[{idx}] {text}", size=9)


def create_architecture_figure() -> None:
    FIG.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(img)

    try:
        title_font = ImageFont.truetype("timesbd.ttf", 42)
        head_font = ImageFont.truetype("timesbd.ttf", 30)
        body_font = ImageFont.truetype("times.ttf", 24)
        small_font = ImageFont.truetype("times.ttf", 20)
    except OSError:
        title_font = ImageFont.load_default()
        head_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    def box(x1, y1, x2, y2, fill, title, items):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=25, outline="black", width=3, fill=fill)
        draw.text(((x1 + x2) / 2, y1 + 28), title, anchor="ma", font=head_font, fill="black")
        y = y1 + 95
        for item in items:
            draw.text((x1 + 25, y), f"• {item}", font=body_font, fill="black")
            y += 48

    draw.text((900, 55), "PROJECT-X System Architecture", anchor="ma", font=title_font, fill="black")
    box(70, 150, 430, 720, "#dbeafe", "Frontend Layer", [
        "Profile Analysis View",
        "Mock Interview Module",
        "Mentor Chat Interface",
        "University Dashboard",
        "Result Visualization",
    ])
    box(500, 150, 860, 720, "#dcfce7", "Flask API Layer", [
        "/analyze_profile",
        "/mock_facial_interview",
        "/mentor_chat",
        "/download_report",
        "Validation and Response Handling",
    ])
    box(930, 150, 1290, 720, "#fef3c7", "Processing Layer", [
        "Resume Parser",
        "TF-IDF + Logistic Regression",
        "Skill Gap Analyzer",
        "Whisper Transcription",
        "DeepFace Emotion Review",
    ])
    box(1360, 150, 1730, 720, "#fee2e2", "Support Layer", [
        "Python Runtime",
        "Docker/PostgreSQL Setup",
        "ReportLab PDF Export",
        "Browser Media Devices",
        "Configuration Files",
    ])

    for x1, x2 in ((430, 500), (860, 930), (1290, 1360)):
        y = 430
        draw.line((x1 + 10, y, x2 - 20, y), fill="black", width=5)
        draw.polygon([(x2 - 20, y - 12), (x2 - 20, y + 12), (x2, y)], fill="black")

    draw.text((465, 395), "HTTP Requests", anchor="ma", font=small_font, fill="black")
    draw.text((895, 395), "Module Calls", anchor="ma", font=small_font, fill="black")
    draw.text((1325, 395), "Runtime Support", anchor="ma", font=small_font, fill="black")

    img.save(FIG)


def main() -> None:
    create_architecture_figure()

    doc = Document()
    configure_page(doc.sections[0])
    set_normal_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    add_run(
        title,
        "PROJECT-X: An AI-Based Career Counselling, Skill Gap Analysis, and Interview Analytics System for Student Employability Support",
        bold=True,
        size=16,
    )

    authors = doc.add_table(rows=1, cols=3)
    authors.alignment = WD_TABLE_ALIGNMENT.CENTER
    authors.autofit = True
    author_blocks = [
        ("Anant Kaushik", "Department of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
        ("Sandeep Rudola", "Department of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
        ("Vipul Chauhan", "Department of Computer Science and Engineering\nGraphic Era Hill University\nDehradun, India"),
    ]
    for cell, (name, aff) in zip(authors.rows[0].cells, author_blocks):
        set_cell_borderless(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, name + "\n", bold=True, size=10)
        add_run(p, aff, size=9)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(6)
    add_run(note, "Conference Paper Draft Prepared from the PROJECT-X Academic Implementation", italic=True, size=9)

    abs_p = doc.add_paragraph()
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.space_after = Pt(3)
    set_paragraph_borderless(abs_p)
    add_run(abs_p, "Abstract—", bold=True, italic=True, size=10)
    add_run(
        abs_p,
        " PROJECT-X is a full-stack academic prototype designed to support students who struggle to interpret their current skill profile, choose relevant career directions, and evaluate interview readiness in a structured way. The system combines resume text analysis, rule-assisted and machine learning based domain prediction, skill-gap identification, mock interview assessment, and mentor-style feedback inside one browser-based workflow. The backend is built with Flask and integrates lightweight natural language processing, TF-IDF with One-vs-Rest Logistic Regression for career-domain prediction, Whisper for transcription, and DeepFace for facial-emotion review. The frontend, implemented in React, presents analysis results as interpretable recommendations rather than opaque scores. A major contribution of the work is the integration of multiple employability-support functions that are often handled by separate tools: profile analysis, interview analytics, and contextual guidance. The paper discusses the motivation behind the system, its architecture, implementation strategy, and observed prototype behavior. The project is positioned not as a production-grade hiring platform, but as an educational decision-support system that can help students identify missing skills, practice communication, and receive more actionable feedback during placement preparation.",
        italic=True,
        size=10,
    )

    kw_p = doc.add_paragraph()
    kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_p.paragraph_format.space_after = Pt(6)
    add_run(kw_p, "Keywords—", bold=True, italic=True, size=10)
    add_run(
        kw_p,
        "career counselling, skill gap analysis, employability analytics, interview assessment, natural language processing, machine learning, Whisper, DeepFace, academic web application",
        italic=True,
        size=10,
    )

    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_page(sec)
    set_columns(sec, 2)

    add_heading_paragraph(doc, "I. INTRODUCTION")
    add_body_paragraph(
        doc,
        "Students preparing for internships and placements often receive fragmented guidance. Resume websites may highlight keywords, interview platforms may focus only on speaking practice, and mentors may provide broad advice without concrete evidence from the student’s current profile. This fragmentation creates a practical gap: students know they need improvement, but they often do not know where to begin or how their present abilities align with target roles.",
    )
    add_body_paragraph(
        doc,
        "PROJECT-X was developed to address this gap through a single academic prototype that combines resume-driven career analysis, missing-skill identification, interview feedback, and mentor-style guidance. Instead of presenting isolated scores, the system attempts to produce interpretable outputs such as recommended domains, matched versus missing skills, transcript issues, and a consolidated employability estimate. The intent is to support learning and preparation rather than automate recruitment decisions.",
    )
    add_body_paragraph(
        doc,
        "The contribution of the project lies in its integration strategy. Lightweight machine learning and natural language processing techniques are combined with browser-based media capture and multimodal feedback so that students can move through an end-to-end preparation workflow inside one application. The resulting system is academically significant because it demonstrates how different analytical modules can be orchestrated into a usable decision-support interface.",
    )

    add_heading_paragraph(doc, "II. RELATED MOTIVATION AND DESIGN CONTEXT")
    add_body_paragraph(
        doc,
        "Existing educational employability tools usually solve only one part of the readiness problem. Resume analyzers focus on textual matching, interview practice tools emphasize speaking behavior, and generic chat systems do not usually understand the learner’s current gaps. This motivated the design of a unified platform where outputs from one stage feed the next stage. For example, extracted skills influence the gap analysis, the gap analysis informs mentor responses, and interview feedback complements the profile score.",
    )
    add_body_paragraph(
        doc,
        "The design context also reflects practical academic constraints. Large labelled datasets, enterprise infrastructure, and industrial validation pipelines were outside the intended scope of the project. Therefore, the implementation favors interpretable techniques, manageable prototype datasets, and modular engineering choices that can be explained clearly in an undergraduate research setting.",
    )

    add_heading_paragraph(doc, "III. SYSTEM OVERVIEW")
    add_body_paragraph(
        doc,
        "The system follows a layered architecture composed of a React frontend, a Flask API layer, a processing layer for parsing and analytics, and a support layer containing runtime dependencies and report generation utilities. The profile analysis flow begins with resume text entered in the browser. The backend parses the text, predicts likely career domains, compares observed skills against curated requirements, and returns a structured result. The interview flow accepts recorded video, performs transcription and facial-emotion review, and generates communication-oriented feedback. A mentor module then uses the analytical context to produce follow-up guidance.",
    )

    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.paragraph_format.space_after = Pt(2)
    run = fig_p.add_run()
    run.add_picture(str(FIG), width=Inches(3.1))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    add_run(cap, "Fig. 1. High-level architecture of the PROJECT-X system.", italic=False, size=8)

    add_heading_paragraph(doc, "IV. METHODOLOGY")
    add_body_paragraph(
        doc,
        "The methodology combines deterministic logic with machine learning so that the system remains both usable and explainable. Resume text is first processed through a lightweight parser that extracts skill keywords and basic profile cues. Career-domain prediction then uses two complementary strategies: a rule-based shortcut for clearly identifiable profiles and a TF-IDF plus One-vs-Rest Logistic Regression pipeline for broader classification over domain-labelled examples. This hybrid approach improves interpretability and helps the prototype produce meaningful outputs even with limited training data.",
    )
    add_body_paragraph(
        doc,
        "Skill-gap analysis is performed through set comparison between extracted student skills and curated requirement lists for target domains. Interview evaluation follows a multimodal pipeline. Whisper transcribes spoken responses from recorded video, transcript analysis identifies filler words and repetition patterns, and DeepFace samples frames to estimate dominant emotional expression. These outputs are aggregated into communication and facial indicators that contribute to an overall employability-oriented score.",
    )
    add_body_paragraph(
        doc,
        "A final mentor-response stage turns the analytical outputs into action-oriented guidance. Instead of asking students to interpret raw predictions alone, the system frames results as practical advice linked to missing skills, likely domain fit, and interview behavior. This makes the platform more supportive for academic use.",
    )

    add_heading_paragraph(doc, "V. IMPLEMENTATION DETAILS")
    add_body_paragraph(
        doc,
        "The frontend was implemented in React to provide a simple multi-view interface for profile analysis, mock interview recording, mentor interaction, and dashboard-style summaries. The backend was implemented in Flask because it offered a straightforward way to expose REST endpoints while coordinating NLP, machine learning, and report-generation modules. Supporting libraries include scikit-learn for text classification, spaCy and regular expressions for resume parsing cues, Whisper for transcription, DeepFace for frame-level emotion analysis, TextBlob for lightweight sentiment support, and ReportLab for PDF report generation.",
    )
    add_body_paragraph(
        doc,
        "The implementation was built iteratively. Profile analysis was completed first because it defines the central employability problem. Interview analytics was then added as a second layer of evaluation, followed by mentor guidance and downloadable reporting. This order helped preserve a usable prototype at each development stage rather than delaying feedback until the end.",
    )

    add_heading_paragraph(doc, "VI. PROTOTYPE OUTPUTS AND OBSERVATIONS")
    add_body_paragraph(
        doc,
        "Prototype-level testing showed that the system can complete the full workflow from profile input to recommendation and interview feedback inside one application. The strongest practical outcome was not a single predictive metric, but the coherence of the outputs. Students receive a recommended domain, a list of missing skills, transcript warnings, dominant facial-emotion cues, and contextual mentor guidance without moving across separate tools.",
    )
    add_body_paragraph(
        doc,
        "The project also revealed several engineering trade-offs. Lightweight models and curated rules make the prototype easier to justify and debug, but they do not provide industrial-grade generalization. Similarly, facial feedback can be useful as a reflective cue, yet it should be interpreted carefully because emotional signals are context-sensitive and not equivalent to objective performance. These observations are important because they shape how such systems should be positioned in educational settings.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Module"
    hdr[1].text = "Technique Used"
    hdr[2].text = "Observed Role in Prototype"
    rows = [
        ("Profile Analysis", "Resume parsing + TF-IDF classification", "Generated domain recommendation and skill context"),
        ("Skill Gap Review", "Set-based comparison of required and extracted skills", "Produced actionable missing-skill output"),
        ("Interview Analytics", "Whisper + DeepFace + transcript issue analysis", "Provided communication and facial feedback"),
        ("Mentor Support", "Prompt-based contextual guidance", "Turned analysis into next-step advice"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.paragraph_format.space_after = Pt(4)
    add_run(cap2, "Table I. Major modules and their practical role in PROJECT-X.", size=8)

    add_heading_paragraph(doc, "VII. LIMITATIONS AND FUTURE WORK")
    add_body_paragraph(
        doc,
        "The present system should be understood as an academic prototype rather than a production-ready decision engine. The classification module relies on limited domain-labelled examples, the interview scoring logic is intentionally lightweight, and long-term storage, authentication, and institutional deployment concerns are only partially addressed. These limitations do not weaken the proof-of-concept value of the project, but they do define the boundary of its current claims.",
    )
    add_body_paragraph(
        doc,
        "Future work can improve the project in three directions: richer and better validated datasets, stronger privacy and consent controls for uploaded media, and broader institutional analytics for batch-level insight. Additional experimentation could also compare model alternatives, improve the calibration of employability scoring, and expand mentor guidance with role-specific learning pathways.",
    )

    add_heading_paragraph(doc, "VIII. CONCLUSION")
    add_body_paragraph(
        doc,
        "PROJECT-X demonstrates that a student-oriented employability support platform can be built by combining interpretable machine learning, natural language processing, multimodal interview analysis, and a clear web interface. Its main value lies in integration: instead of offering fragmented assistance, it connects profile understanding, skill-gap identification, interview reflection, and guidance in one workflow. For an academic project, this makes the system both technically meaningful and educationally relevant.",
    )

    add_heading_paragraph(doc, "ACKNOWLEDGMENT")
    add_body_paragraph(
        doc,
        "The authors acknowledge the guidance of the Department of Computer Science and Engineering, Graphic Era Hill University, and the support received during the development and documentation of this undergraduate project.",
        first_line_cm=0.0,
    )

    add_heading_paragraph(doc, "REFERENCES")
    references = [
        "C. D. Manning, P. Raghavan, and H. Schutze, Introduction to Information Retrieval. Cambridge, U.K.: Cambridge Univ. Press, 2008.",
        "T. M. Mitchell, Machine Learning. New York, NY, USA: McGraw-Hill, 1997.",
        "D. Jurafsky and J. H. Martin, Speech and Language Processing, 3rd ed. draft. Stanford, CA, USA: Stanford Univ., 2025.",
        "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” J. Mach. Learn. Res., vol. 12, pp. 2825–2830, 2011.",
        "The Pallets Projects, “Flask documentation,” 2026. [Online]. Available: https://flask.palletsprojects.com/",
        "A. Geron, Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow, 2nd ed. Sebastopol, CA, USA: O’Reilly, 2019.",
        "A. Radford et al., “Robust speech recognition via large-scale weak supervision,” in Proc. 40th Int. Conf. Mach. Learn., 2023.",
        "O. M. Parkhi, A. Vedaldi, and A. Zisserman, “Deep face recognition,” in Proc. British Machine Vision Conf., 2015.",
        "S. Bird, E. Klein, and E. Loper, Natural Language Processing with Python. Sebastopol, CA, USA: O’Reilly, 2009.",
    ]
    for i, ref in enumerate(references, 1):
        add_reference(doc, i, ref)

    doc.save(OUT)
    print(f"Saved conference paper to: {OUT}")


if __name__ == "__main__":
    main()
